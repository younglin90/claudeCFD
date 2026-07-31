#!/usr/bin/env python3
"""Convert 1d_euler_imex_method_revised.md → DOCX.

Pure-Python Markdown → DOCX converter tailored for our manuscript.
Handles headings, paragraphs, bullet lists, tables (pipe syntax),
fenced code blocks, blockquotes, inline bold (**) and inline code (`).
Math is left as plain text (rendered by reader).
"""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path('/home/younglin90/work/claude_code/claudeCFD')
import sys
SRC_NAME = sys.argv[1] if len(sys.argv) > 1 else '1d_euler_imex_method_revised.md'
SRC = ROOT / 'docs' / 'manuscripts' / SRC_NAME
DST = SRC.with_suffix('.docx')


def add_inline_runs(paragraph, text):
    """Parse **bold**, *italic*, `code` inside a line and add runs."""
    # Tokenize.  Order: code (`...`), bold (**...**), italic (*...*).
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('`') and token.endswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            try:
                r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Consolas')
            except Exception:
                pass
            r.font.size = Pt(9)
        elif token.startswith('**') and token.endswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('*') and token.endswith('*') and len(token) > 2:
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        else:
            paragraph.add_run(token)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_table_from_lines(doc, header_line, sep_line, body_lines):
    """Pipe-table: | a | b | c |"""
    def split_cells(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    headers = split_cells(header_line)
    n = len(headers)
    rows = [split_cells(r) for r in body_lines if r.strip()]

    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        if i < len(hdr):
            cell = hdr[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9)
    for ri, row in enumerate(rows, start=1):
        for ci in range(n):
            val = row[ci] if ci < len(row) else ''
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            add_inline_runs(p, val)
            for r in p.runs:
                r.font.size = Pt(9)


def main():
    text = SRC.read_text(encoding='utf-8')
    lines = text.splitlines()

    doc = Document()
    # Default body style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Tighten section margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    in_code = False
    code_buf: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if line.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                # close
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                r = p.add_run('\n'.join(code_buf))
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                in_code = False
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r'-{3,}', line.strip()):
            p = doc.add_paragraph()
            p.add_run('—' * 30).italic = True
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            heading_level = min(level, 4)  # docx supports Heading 1-9 generally
            h = doc.add_heading(level=heading_level)
            add_inline_runs(h, heading_text)
            i += 1
            continue

        # Tables (pipe-style)
        if line.lstrip().startswith('|') and i + 1 < len(lines) \
                and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i + 1]) \
                and '|' in lines[i + 1]:
            header = line
            sep = lines[i + 1]
            body: list[str] = []
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                body.append(lines[i])
                i += 1
            add_table_from_lines(doc, header, sep, body)
            continue

        # Bullet lists
        m = re.match(r'^\s*[-*]\s+(.*)$', line)
        if m:
            content = m.group(1)
            p = doc.add_paragraph(style='List Bullet')
            add_inline_runs(p, content)
            i += 1
            continue

        # Numbered lists
        m = re.match(r'^\s*\d+\.\s+(.*)$', line)
        if m:
            content = m.group(1)
            p = doc.add_paragraph(style='List Number')
            add_inline_runs(p, content)
            i += 1
            continue

        # Blockquote
        if line.lstrip().startswith('>'):
            content = line.lstrip()[1:].lstrip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run('  ' + content)
            r.italic = True
            i += 1
            continue

        # Empty
        if not line.strip():
            i += 1
            continue

        # Math display lines: $$ ... $$
        if line.strip().startswith('$$') and line.strip().endswith('$$') and len(line.strip()) > 4:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run(line.strip().strip('$').strip())
            r.font.name = 'Cambria Math'
            r.font.italic = True
            i += 1
            continue
        if line.strip() == '$$':
            # Multi-line display math
            i += 1
            buf: list[str] = []
            while i < len(lines) and lines[i].strip() != '$$':
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing $$
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run('\n'.join(buf))
            r.font.name = 'Cambria Math'
            r.font.italic = True
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1

    doc.save(DST)
    size_kb = DST.stat().st_size / 1024
    print(f'wrote {DST} ({size_kb:.1f} KB)')


if __name__ == '__main__':
    main()

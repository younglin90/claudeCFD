from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화/"
    "SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d.docx"
)
OUT = SRC.with_name("SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_cavitydiag.docx")


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def main() -> None:
    doc = Document(SRC)
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("D2Q9 중심 검증."):
            set_text(
                p,
                "D2Q9 중심 검증. 본문의 정량 결과와 확장 검증은 모두 2D D2Q9 격자에 기반한다. "
                "따라서 본 논문은 2D 정상상태 LBM residual 가속의 가능성과 한계를 다루며, "
                "이 범위를 넘어서는 검증은 본 원고의 claim에 포함하지 않는다.",
            )
        elif text.startswith("대규모 격자"):
            set_text(
                p,
                "대규모 격자 scaling 시험. FFT-기반 전처리기의 통신 비용과 가속 효과의 "
                "break-even point를 측정하고, 병렬 FFT와의 결합 가능성을 평가한다.",
            )
        elif text.startswith("정량 해석. Kolmogorov는 N=128/256"):
            set_text(
                p,
                text
                + " 추가 진단으로 Re=400, N=65 Safe-NN 종료해를 Picard-LBM으로 후처리한 결과, "
                "residual을 4.63e-7에서 2.95e-8로 낮추면 tight-baseline 대비 relative L2가 "
                "1.51e-2에서 1.97e-3으로 감소했다. 따라서 cavity 그림에서 보이는 잔류 수치 "
                "진동은 주로 5e-7 종료 기준에서 남은 kinetic/high-frequency 모드의 미감쇠로 "
                "판단된다.",
            )
        elif text.startswith("Stiff regime의 가속 저하."):
            set_text(
                p,
                "Stiff regime의 가속 저하와 잔류 진동. Cavity Re=400에서 LBE-call 가속률은 "
                "5.66x-6.71x 범위로 떨어지고, 현재 Python/JFNK 구현의 wall-clock은 baseline보다 "
                "느리다. Re=400 N=65 진단 계산에서는 Safe-NN 종료 residual 4.63e-7의 해를 "
                "Picard post-relaxation으로 2.95e-8까지 낮추자 baseline 대비 relative L2가 "
                "1.51e-2에서 1.97e-3으로 감소했다. 이는 그림의 작은 진동이 물리적 비정상성이 "
                "아니라 loose cavity tolerance와 post-relaxation 부족에서 남은 비평형 모드임을 "
                "시사한다. Re=1000(N=129)에서는 line-search safeguard가 NaN 발산을 막아 residual "
                "수렴은 달성했지만 LBE-call gain은 1.17x에 그쳤고 Ghia deviation도 크다.",
            )
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

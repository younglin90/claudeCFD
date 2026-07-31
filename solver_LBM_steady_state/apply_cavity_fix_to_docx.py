from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches


SRC = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화/"
    "SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_cavitydiag.docx"
)
OUT = SRC.with_name("SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_cavityfixed.docx")
DATA = Path("paper_revision_data/cavity_polished_results.json")


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def add_picture_after(paragraph, image_path: str, caption: str) -> None:
    body = paragraph._element.getparent()
    idx = body.index(paragraph._element)
    fig_p = paragraph._parent.add_paragraph()
    fig_p.add_run().add_picture(image_path, width=Inches(6.4))
    cap_p = paragraph._parent.add_paragraph(caption)
    body.remove(fig_p._element)
    body.remove(cap_p._element)
    body.insert(idx + 1, fig_p._element)
    body.insert(idx + 2, cap_p._element)


def main() -> None:
    data = json.loads(DATA.read_text(encoding="utf-8"))
    doc = Document(SRC)

    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("정량 해석. Kolmogorov는 N=128/256"):
            set_text(
                p,
                "정량 해석. Kolmogorov는 N=128/256 모두 143x/422x LBE-call scaling을 유지하며 "
                "1e-9 tolerance에 도달했다(Figure R1). Channel은 LBE-call 자체는 "
                "17.6x/50.5x 감소했지만 300 외부 반복 내에 1e-9 tolerance에 도달하지 못해 "
                "residual plateau case로 분류했다. Re=400(N=65)는 loose cavity tolerance "
                "5e-7에서 LBE-call 6.71x를 보였지만, 그림의 잔류 진동을 제거하기 위해 "
                "Safe-NN 종료 후 Picard final-polish를 5e-8까지 추가했다. Polished 결과는 "
                "Safe residual 4.53e-8, tight-baseline 대비 relative L2 2.88e-3, Ghia "
                "centerline max deviation 5.45e-2이며, tight-tolerance 기준 LBE-call 가속은 "
                "2.38x이다. Re=1000은 N=65에서 overflow/NaN이 발생했으므로 N=129 "
                "line-search safeguard로 재계산했으며, residual은 4.99e-7까지 감소했지만 "
                "LBE-call gain은 1.17x, wall-clock은 0.04x, Ghia deviation은 3.29e-1이다. "
                "이 행은 고-Re BGK/전처리 한계 자료이지 가속 성공 claim이 아니다.",
            )
            add_picture_after(
                p,
                data["figure"],
                "Figure R10. Cavity Re=400, N=65 after final Picard polish. The Safe-NN field is "
                "post-relaxed to residual 4.53e-8; the relative L2 velocity difference against the "
                "tight Picard baseline is 2.88e-3.",
            )
        elif text.startswith("종합. 핵심 6 검증 사례"):
            set_text(
                p,
                "종합. 핵심 6 검증 사례에서는 Safe-NN-SCMK가 모두 지정 tolerance 내에서 "
                "수렴했다. 전체 6 case 기준 LBE-call 가속은 산술평균 45.1x / 기하평균 19.2x이며, "
                "Re=400 stress test를 제외한 5개 표준 case 산술평균은 53.0x이다. 확장 검증에서는 "
                "N=128/256 Kolmogorov, Re=400(N=65) polished cavity, backward-step, cylinder "
                "analogue가 추가 수렴 근거를 제공했다. 반면 N=128/256 channel의 1e-9 plateau, "
                "Re=1000 cavity의 약한 정확도/느린 wall-clock, T-junction의 큰 relative L2는 "
                "현재 알고리즘과 Python 구현의 한계로 본문에 명시한다.",
            )
        elif text.startswith("Stiff regime의 가속 저하와 잔류 진동."):
            set_text(
                p,
                "Stiff regime의 가속 저하와 final-polish 보정. Cavity Re=400에서 loose "
                "tolerance 기준 LBE-call 가속률은 5.66x-6.71x 범위이지만, 이 상태의 해에는 "
                "kinetic/high-frequency 비평형 모드가 잔류해 그림에서 작은 수치 진동이 보일 수 "
                "있다. 이를 해결하기 위해 cavity 결과에는 Safe-NN 종료 후 Picard final-polish를 "
                "적용한다. Re=400 N=65에서는 residual이 4.63e-7에서 4.53e-8로 낮아지고, "
                "tight-baseline 대비 relative L2가 1.51e-2에서 2.88e-3으로 감소했다. 다만 "
                "tight-tolerance 기준 LBE-call speedup은 2.38x로 낮아지며 wall-clock은 현재 "
                "Python/JFNK 구현에서 baseline보다 빠르지 않다. Re=1000(N=129)에서는 line-search "
                "safeguard가 NaN 발산을 막아 residual 수렴은 달성했지만 LBE-call gain은 1.17x에 "
                "그쳤고 Ghia deviation도 크다.",
            )
        elif text.startswith("본 연구는 표준 LBM의 단순성과 메모리 footprint"):
            set_text(
                p,
                "본 연구는 표준 LBM의 단순성과 메모리 footprint를 보존한 채 정상상태 가속을 "
                "도입하는 framework를 제안하고, 핵심 6-case 및 확장 2D 검증을 통해 효과와 "
                "한계를 정량적으로 평가했다. 전체 6 case 기준 산술평균/기하평균 LBE-call 가속은 "
                "45.1x/19.2x이며, 5개 표준 case 산술평균은 53.0x이다. 추가 계산은 Kolmogorov "
                "N=128/256에서 강한 scaling 이득을 확인했고, Re=400 cavity는 final-polish로 "
                "잔류 수치 진동을 제거할 수 있음을 보였다. 동시에 channel 1e-9 plateau, "
                "Re=1000 cavity의 약한 정확도와 wall-clock 저하, T-junction의 상대오차 증폭을 "
                "드러냈다. 따라서 현재 원고의 결론은 'native LBM operator를 유지한 residual "
                "가속 가능성'과 'stiff/high-Re regime에서 필요한 후속 보완'을 함께 주장하는 "
                "형태가 가장 타당하다.",
            )

    # Update the expanded cavity table row if present.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.startswith("Cavity Re=400 N=65") and len(row.cells) >= 10:
                vals = [
                    "Cavity Re=400 N=65 polished",
                    f"{data['nu']:.5f}",
                    f"{data['omega']:.5f}",
                    f"{data['baseline_lbe_tight']:,}",
                    f"{data['safe_total_lbe']:,}",
                    f"{data['speedup_lbe_tight']:.2f}",
                    f"{data['speedup_wall_tight']:.2f}",
                    f"{data['safe_final_residual']:.3e}",
                    f"{data['safe_ghia']['centerline_max']:.3e}",
                    "True",
                ]
                for cell, val in zip(row.cells, vals):
                    cell.text = val
            elif row_text.startswith("Cavity Re=400 N=65") and "Newton steps" not in row_text:
                stats = data["safe_stats"]
                vals = [
                    "Cavity Re=400 N=65 polished",
                    str(stats["newton_steps"]),
                    str(stats["lookahead_evaluations"]),
                    str(stats["lookahead_rejections"]),
                    str(stats["residual_increase_restarts"]),
                    str(stats["nan_fallbacks"]),
                    str(stats["short_K_steps"]),
                    str(stats.get("line_search_rejections", 0)),
                ]
                for cell, val in zip(row.cells, vals):
                    cell.text = val

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

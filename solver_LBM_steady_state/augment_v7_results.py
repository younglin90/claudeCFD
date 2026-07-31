from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


SRC = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화/"
    "SafeNN_LBM_Paper_V7_KR_final.docx"
)
OUT = SRC.with_name("SafeNN_LBM_Paper_V7_KR_final_results_augmented.docx")


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_before(paragraph, text: str, style: str | None = None):
    new_p = paragraph.insert_paragraph_before(text)
    if style:
        try:
            new_p.style = style
        except KeyError:
            pass
    return new_p


def find_paragraph(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise RuntimeError(f"paragraph not found: {prefix}")


def find_paragraph_optional(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def append_case_rows(table) -> None:
    if any("Kolmogorov scaling N=64" in row.cells[1].text for row in table.rows):
        return

    rows = [
        (
            "7",
            "Kolmogorov scaling N=64",
            "64x64 periodic",
            "omega=1.85",
            "Body force, U0=0.04",
            "1e-7",
        ),
        (
            "8",
            "Channel scaling N=64",
            "64x64",
            "omega=1.85",
            "Body force, no-slip walls",
            "1e-7",
        ),
        (
            "9",
            "Kolmogorov strict scaling N=128/256",
            "128x128, 256x256 periodic",
            "nu=0.05 (omega approx. 1.538)",
            "Body force scaled for Uamp=0.05",
            "1e-9",
        ),
        (
            "10",
            "Channel strict scaling N=128/256",
            "128x128, 256x256",
            "nu=0.05 (omega approx. 1.538)",
            "Body force scaled for Umax=0.05, no-slip walls",
            "1e-9; Safe-NN plateau rows reported separately",
        ),
        (
            "11",
            "Lid-driven cavity Re=400 refinement",
            "65x65",
            "Re=400, nu=0.016, omega=1.8248",
            "U_lid=0.1, no-slip",
            "5e-7",
        ),
        (
            "12",
            "Lid-driven cavity Re=1000 high-Re stress",
            "129x129",
            "Re=1000, nu=0.0128, omega=1.8574",
            "U_lid=0.1, no-slip; line-search safeguard",
            "5e-7; limitation case",
        ),
        (
            "13",
            "Backward-facing step mask",
            "64x64 voxel mask",
            "nu=0.05",
            "Fx=1.5e-5, bounce-back mask",
            "1e-7",
        ),
        (
            "14",
            "Cylinder wake analogue",
            "64x64 voxel mask",
            "nu=0.04",
            "Fx=1.0e-5, cylinder radius=6",
            "1e-7; steady analogue, not unsteady shedding",
        ),
        (
            "15",
            "T-junction mask",
            "64x64 voxel mask",
            "nu=0.05",
            "Fx=8e-6 plus branch Fy=-8e-6",
            "1e-7; convergence stress case",
        ),
        (
            "16",
            "D3Q19 smoke tests",
            "16x16x16",
            "D3Q19 BGK",
            "Kolmogorov and channel smoke checks only",
            "1e-7; not production 3D validation",
        ),
    ]
    for row_values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, row_values):
            cell.text = value


def main() -> None:
    doc = Document(SRC)

    # Results overview and complete validation case list.
    set_paragraph_text(
        find_paragraph(doc, "본 장은 baseline LBM"),
        "본 장은 baseline Picard-LBM, Anderson-accelerated LBM, 그리고 제안 알고리즘 "
        "Safe-NN-SCMK의 비교 결과를 제시한다. §3.1에서는 핵심 6-case의 수렴성"
        "(LBE-call 수, wall-clock, 수렴 이력)을, §3.2에서는 정확도(RMS, L_inf, "
        "relative L2, Ghia centerline deviation, mass drift)를 정량 검증한다. 이어서 "
        "§3.3에서는 추가 검증 case 전체 목록을 기준으로 N=64/128/256 scaling, "
        "cavity Re=400/1000 stress test, backward-facing step, cylinder-wake analogue, "
        "T-junction mask, 그리고 별도 D3Q19 smoke test를 성공/plateau/한계 case로 "
        "구분하여 보고한다.",
    )
    set_paragraph_text(
        find_paragraph(doc, "Table 1. 검증 6-case"),
        "Table 1. 본문 및 확장 검증 전체 case 설정. Case 1-6은 핵심 정량 비교 "
        "benchmark이고, Case 7-16은 major-revision 보완을 위해 추가한 scaling, "
        "고-Re stress, mask-flow, smoke-test 검증이다.",
    )
    append_case_rows(doc.tables[1])
    set_paragraph_text(
        find_paragraph(doc, "Case 1–4, 6 의 수렴 기준"),
        "Case 1-6의 수렴 기준은 기본적으로 1e-7이며, cavity case는 5e-7로 설정했다. "
        "추가 scaling 검증 중 N=128/256 Kolmogorov와 channel은 더 엄격한 1e-9로 "
        "재계산하여 residual plateau 여부를 확인했다. Re=1000 cavity는 N=65 BGK "
        "설정에서 overflow/NaN이 발생해, 본문에는 N=129 line-search safeguard 결과만 "
        "보고한다. 모든 case는 동일 계열의 초기 분포에서 출발했으며, 실패 또는 plateau "
        "case는 성공 검증이 아니라 한계 관찰로 분리했다.",
    )

    # Replace ambiguous speedup prose.
    set_paragraph_text(
        find_paragraph(doc, "선형/단일 모드 case"),
        "선형/단일 모드 case (Kolmogorov, Channel, Couette): 22.5x-194.3x의 큰 "
        "LBE-call 가속. Couette의 정상해는 선형 profile이고 저차 Fourier 성분에 "
        "집중되어 있어 AP-Schur 전처리기가 거시 모멘텀을 매우 빠르게 보정한다. 그 결과 "
        "외부 반복 단 1회(총 30 LBE)로 tolerance에 도달했다.",
    )
    set_paragraph_text(
        find_paragraph(doc, "비선형 vortex case"),
        "비선형 vortex/복잡 경계 case (cavity Re=100, multi-cylinder): 6.2x-10.2x의 "
        "중간 LBE-call 가속. 다중 mode와 bounce-back 경계가 존재하므로 smooth "
        "periodic case보다 가속률은 낮지만, safeguard와 JFNK 보정이 결합되어 baseline "
        "수렴해와 동일 tolerance 수준의 해를 얻었다.",
    )
    set_paragraph_text(
        find_paragraph(doc, "stiff regime (cavity Re=400)"),
        "stiff regime (cavity Re=400): 5.66x의 보수적 LBE-call 가속. residual "
        "safeguard가 외삽을 자주 거부하여 momentum 기여가 감소하고, AP-Schur "
        "전처리기의 Chapman-Enskog 기반 근사가 강한 streamline 구조에서 약해진다. "
        "현재 Python/JFNK 구현의 wall-clock은 baseline보다 느리므로, 이 case는 "
        "절대 시간 가속 주장이 아니라 stiff regime에서의 안정 수렴 stress test로 "
        "해석한다.",
    )

    # Add a visible subsection heading for the extended validations.
    p_extra = find_paragraph(doc, "위 6 case 외에, mask 기반")
    insert_paragraph_before(p_extra, "3.3 확장 검증, 수치 진동, 한계 case", "Heading 2")
    set_paragraph_text(
        p_extra,
        "위 6 case 외에 Table 1의 Case 7-16을 추가 검증했다. 목적은 논문의 평균 "
        "가속률을 높이는 것이 아니라, reviewer가 요구할 가능성이 높은 격자 scaling, "
        "고-Re cavity, 복잡 mask geometry, 그리고 3D 이식 가능성의 경계를 분리해서 "
        "보여주는 것이다. 따라서 아래 결과는 성공 case와 plateau/limitation case를 "
        "명시적으로 구분한다.",
    )
    set_paragraph_text(
        find_paragraph(doc, "정량 해석. Figure R5-R6"),
        "정량 해석. Figure R5-R6은 backward-facing step에서 Safe-NN이 Picard 대비 "
        "16.2x LBE-call 감소와 8.13x wall-clock 가속을 얻고, cylinder-wake analogue는 "
        "2.43x LBE-call / 1.46x wall-clock으로 장애물 지배 흐름에서 이득이 제한됨을 "
        "보여준다. T-junction은 residual 기준으로는 31 LBE 만에 수렴해 32.3x "
        "LBE-call 감소를 보였지만, 평균 기준 속도가 매우 작아 relative L2가 0.609로 "
        "증폭된다. 따라서 T-junction은 정확도 benchmark가 아니라 좁은 mask network의 "
        "잔차 수렴 stress case로만 사용한다. Cylinder case 역시 정상 주기 mask analogue이며 "
        "unsteady vortex shedding 검증으로 해석하지 않는다.",
    )
    set_paragraph_text(
        find_paragraph(doc, "정량 해석. Kolmogorov는 N=128"),
        "정량 해석. Kolmogorov는 N=128/256 모두 143x/422x LBE-call scaling을 유지하며 "
        "1e-9 tolerance에 도달했다(Figure R1). Channel은 LBE-call 자체는 "
        "17.6x/50.5x 감소했지만 300 외부 반복 내에 1e-9 tolerance에 도달하지 못해 "
        "residual plateau case로 분류했다. Re=400(N=65)는 LBE-call 6.71x와 Ghia "
        "centerline max deviation 5.22e-2로 기존 Re=400 stress 관찰을 재현했다. "
        "Re=1000은 N=65에서 overflow/NaN이 발생했으므로 N=129 line-search safeguard로 "
        "재계산했으며, residual은 4.99e-7까지 감소했지만 LBE-call gain은 1.17x, "
        "wall-clock은 0.04x, Ghia deviation은 3.29e-1이다. 이 행은 고-Re BGK/전처리 "
        "한계 자료이지 가속 성공 claim이 아니다.",
    )
    set_paragraph_text(
        find_paragraph(doc, "종합. 모든 6 검증 사례"),
        "종합. 핵심 6 검증 사례에서는 Safe-NN-SCMK가 모두 지정 tolerance 내에서 "
        "수렴했다. 전체 6 case 기준 LBE-call 가속은 산술평균 45.1x / 기하평균 19.2x이며, "
        "Re=400 stress test를 제외한 5개 표준 case 산술평균은 53.0x이다. 확장 검증에서는 "
        "N=128/256 Kolmogorov와 Re=400(N=65), backward-step, cylinder analogue가 "
        "추가 수렴 근거를 제공했다. 반면 N=128/256 channel의 1e-9 plateau, Re=1000 "
        "cavity의 약한 정확도/느린 wall-clock, T-junction의 큰 relative L2는 현재 "
        "알고리즘과 Python 구현의 한계로 본문에 명시한다.",
    )

    # Discussion/conclusion consistency after the additional calculations.
    set_paragraph_text(
        find_paragraph(doc, "Stiff regime의 가속 저하."),
        "Stiff regime의 가속 저하. Cavity Re=400에서 LBE-call 가속률은 5.66x-6.71x "
        "범위로 떨어지고, 현재 Python/JFNK 구현의 wall-clock은 baseline보다 느리다. "
        "Re=1000(N=129)에서는 line-search safeguard가 NaN 발산을 막아 residual 수렴은 "
        "달성했지만 LBE-call gain은 1.17x에 그쳤고 Ghia deviation도 크다. 이는 "
        "(a) safeguard가 외삽을 빈번히 거부하고, (b) BGK 고-Re 설정에서 AP-Schur "
        "전처리기의 선형화 근사가 약해지며, (c) finite-difference JVP와 line-search가 "
        "wall-clock 비용을 증가시키기 때문이다.",
    )
    set_paragraph_text(
        find_paragraph(doc, "Tolerance 의존성."),
        "Tolerance 의존성. 본문 핵심 실험의 nonlinear tolerance는 1e-7(cavity 5e-7)이다. "
        "더 엄격한 1e-9 조건에서는 Kolmogorov N=128/256은 수렴했지만 channel N=128/256은 "
        "300 외부 반복 내 residual plateau를 보였다(Figure R1). 따라서 본 방법의 강점은 "
        "production tolerance에서의 정상해 도달 가속이며, 1e-9 이하의 sensitivity "
        "analysis에는 추가 preconditioner 또는 MRT 확장이 필요하다.",
    )
    set_paragraph_text(
        find_paragraph(doc, "본 연구는 표준 LBM의 단순성과 메모리 footprint"),
        "본 연구는 표준 LBM의 단순성과 메모리 footprint를 보존한 채 정상상태 가속을 "
        "도입하는 framework를 제안하고, 핵심 6-case 및 확장 2D 검증을 통해 효과와 "
        "한계를 정량적으로 평가했다. 전체 6 case 기준 산술평균/기하평균 LBE-call 가속은 "
        "45.1x/19.2x이며, 5개 표준 case 산술평균은 53.0x이다. 추가 계산은 Kolmogorov "
        "N=128/256에서 강한 scaling 이득을 확인하는 동시에, channel 1e-9 plateau, "
        "Re=1000 cavity의 약한 정확도와 wall-clock 저하, T-junction의 상대오차 증폭을 "
        "드러냈다. 따라서 현재 원고의 결론은 'native LBM operator를 유지한 residual "
        "가속 가능성'과 'stiff/high-Re regime에서 필요한 후속 보완'을 함께 주장하는 "
        "형태가 가장 타당하다.",
    )

    # Remove author-facing revision notes that should not remain after conclusion.
    start = find_paragraph_optional(doc, "Major-revision 보완 계산 요약")
    end = find_paragraph_optional(doc, "References")
    if start is not None and end is not None:
        deleting = False
        for p in list(doc.paragraphs):
            if p._element is start._element:
                deleting = True
            if p._element is end._element:
                deleting = False
            if deleting:
                delete_paragraph(p)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

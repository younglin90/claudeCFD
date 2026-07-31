#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build Results_section_KO.docx from results_ko.md using python-docx.
Handles: headings, blockquotes, images (with {width=Ncm}), pipe tables,
**bold**, inline $math$ (rendered to unicode), and \\newpage."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
MD = HERE / "results_ko.md"
OUT = HERE / "Results_section_KO.docx"

LATIN = "Times New Roman"
EAST = "Malgun Gothic"
CONTENT_W = 16.3  # cm, Letter w/ 1in margins

MATH = [
    (r"\times", "×"), (r"\le", "≤"), (r"\ge", "≥"), (r"\approx", "≈"),
    (r"\alpha", "α"), (r"\tau", "τ"), (r"\delta", "δ"), (r"\rho", "ρ"),
    (r"\Omega", "Ω"), (r"\in", "∈"), (r"\to", "→"), (r"\|", "‖"),
    (r"\,", " "), (r"\tfrac", ""), (r"\mathbf", ""), (r"\mathcal", ""),
    (r"\text", ""),
]
SUP = {"-": "⁻", "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
       "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹"}


def mathify(s: str) -> str:
    # superscripts like ^{-6} or ^2 -> unicode
    def sup(m):
        body = m.group(1) or m.group(2)
        if all(c in SUP for c in body):
            return "".join(SUP[c] for c in body)
        return "^" + body
    s = re.sub(r"\^\{([^}]*)\}", lambda m: sup(re.match(r"()(.*)", m.group(1))), s)
    s = re.sub(r"\^\{([^}]*)\}|\^(\S)", sup, s)
    for a, b in MATH:
        s = s.replace(a, b)
    s = s.replace("{", "").replace("}", "")
    s = re.sub(r"\\([a-zA-Z]+)", r"\1", s)  # drop leftover commands
    s = s.replace("\\", "")
    return s


def strip_inline_math(text: str) -> str:
    # replace $...$ spans with mathified content
    return re.sub(r"\$([^$]*)\$", lambda m: mathify(m.group(1)), text)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = LATIN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), EAST)
    rFonts.set(qn("w:ascii"), LATIN)
    rFonts.set(qn("w:hAnsi"), LATIN)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_inline(par, text, base_size=10.5, base_bold=False, color=None):
    """Add text with **bold** and $math$ handling."""
    text = strip_inline_math(text)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            r = par.add_run(p[2:-2])
            set_run_font(r, base_size, True, color)
        else:
            r = par.add_run(p)
            set_run_font(r, base_size, base_bold, color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), fill)
    tcPr.append(sh)


def add_table(doc, rows):
    header = rows[0]
    body = rows[1:]
    t = doc.add_table(rows=len(rows), cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].clear()
        add_inline(cell.paragraphs[0], h, base_size=9, base_bold=True)
        shade_cell(cell, "D9E2F3")
    for i, row in enumerate(body, 1):
        for j, val in enumerate(row):
            if j >= len(header):
                continue
            cell = t.rows[i].cells[j]
            cell.paragraphs[0].clear()
            add_inline(cell.paragraphs[0], val, base_size=9)
    doc.add_paragraph()


def add_image(doc, caption, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img = HERE / path
    if img.exists():
        w = min(width_cm, CONTENT_W)
        p.add_run().add_picture(str(img), width=Cm(w))
    else:
        add_inline(p, f"[missing image: {path}]")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(cap, caption, base_size=9)
    cap.paragraph_format.space_after = Pt(8)


def main():
    lines = MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    # base style
    st = doc.styles["Normal"]
    st.font.name = LATIN
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), EAST)
    # page size Letter + 1in margins
    sec = doc.sections[0]
    from docx.shared import Inches
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1.0))

    i = 0
    n = len(lines)
    img_re = re.compile(r"!\[(.*)\]\((.*?)\)(?:\{width=([\d.]+)cm\})?")
    while i < n:
        line = lines[i].rstrip("\n")
        s = line.strip()
        if not s:
            i += 1
            continue
        # page break
        if s == "\\newpage":
            doc.add_page_break()
            i += 1
            continue
        # headings
        m = re.match(r"^(#{1,3})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 3))
            h.clear()
            add_inline(h, m.group(2),
                       base_size={1: 15, 2: 13, 3: 11.5}[level], base_bold=True,
                       color=RGBColor(0x1F, 0x1F, 0x1F))
            i += 1
            continue
        # image
        mi = img_re.match(s)
        if mi:
            cap, path, w = mi.group(1), mi.group(2), mi.group(3)
            add_image(doc, cap, path, float(w) if w else 12.0)
            i += 1
            continue
        # blockquote
        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip()[1:].strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.4)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, " ".join(buf), base_size=9.5, color=RGBColor(0x40, 0x40, 0x40))
            # left border bar
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
            left.set(qn("w:space"), "8"); left.set(qn("w:color"), "9DB7D6")
            pbdr.append(left); pPr.append(pbdr)
            continue
        # table
        if s.startswith("|"):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            parsed = []
            for r, rowline in enumerate(tbl):
                cells = [c.strip() for c in rowline.strip().strip("|").split("|")]
                if re.match(r"^[\s:|-]+$", "|".join(cells)):
                    continue  # separator
                parsed.append(cells)
            if parsed:
                add_table(doc, parsed)
            continue
        # normal paragraph (gather until blank)
        buf = [s]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^(#{1,3}\s|>|\||!\[|\\newpage)", lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, " ".join(buf))

    doc.save(str(OUT))
    print("saved", OUT)


if __name__ == "__main__":
    main()

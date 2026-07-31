"""Append field/convergence figures to the figure-enhanced DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches


ONEDRIVE_DIR = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화"
)
SOURCE_DOCX = (
    ONEDRIVE_DIR
    / "SafeNN_LBM_Paper_V4_KR_v2_expanded_major_revision_refs_extra_benchmarks_figures.docx"
)
OUTPUT_DOCX = (
    ONEDRIVE_DIR
    / "SafeNN_LBM_Paper_V4_KR_v2_expanded_major_revision_refs_extra_benchmarks_all_figures.docx"
)
FIG_DIR = Path("paper_revision_data/figures")


FIGURES = [
    (
        FIG_DIR / "fig_extra_benchmark_convergence.png",
        "Figure R5. Native residual histories for the additional 2D benchmarks. Solid lines denote Safe-NN and dashed lines denote Picard-LBM.",
        6.5,
    ),
    (
        FIG_DIR / "fig_extra_benchmark_lbe_to_tolerance.png",
        "Figure R6. LBE-call cost to reach tolerance for additional 2D benchmarks.",
        6.2,
    ),
    (
        FIG_DIR / "fig_field_backward_step.png",
        "Figure R7. Backward-facing step analogue: Picard velocity magnitude, Safe-NN velocity magnitude, and absolute difference.",
        6.7,
    ),
    (
        FIG_DIR / "fig_field_cylinder_wake.png",
        "Figure R8. Cylinder wake analogue: Picard velocity magnitude, Safe-NN velocity magnitude, and absolute difference.",
        6.7,
    ),
    (
        FIG_DIR / "fig_field_t_junction.png",
        "Figure R9. T-junction analogue: Picard velocity magnitude, Safe-NN velocity magnitude, and absolute difference.",
        6.7,
    ),
]


def add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.bold = True


def main() -> None:
    document = Document(SOURCE_DOCX)
    add_heading(document, "Additional field and convergence figures")
    document.add_paragraph(
        "These figures show the actual velocity-magnitude fields and residual histories for the added 2D benchmark cases."
    )
    for path, caption, width in FIGURES:
        if not path.exists():
            raise FileNotFoundError(path)
        document.add_picture(str(path), width=Inches(width))
        document.add_paragraph(caption)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()

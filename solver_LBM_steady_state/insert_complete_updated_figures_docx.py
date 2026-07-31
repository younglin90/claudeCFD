from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches


ONEDRIVE_DIR = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화"
)
SOURCE = ONEDRIVE_DIR / "SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_cavityfixed.docx"
OUTPUT = ONEDRIVE_DIR / "SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_allfigures_updated.docx"
FIG = Path("paper_revision_data/figures_complete")


FIGURES = [
    (
        "Figure U1. Complete convergence histories for the six core 2D benchmarks.",
        FIG / "figU1_core6_convergence_updated.png",
        "Picard-LBM is shown by dashed gray curves and Safe-NN-SCMK by solid red curves. "
        "The Re=400 cavity row is retained as a stiff stress-test; it exhibits a residual floor under the coarse N=49 grid.",
        6.8,
    ),
    (
        "Figure U2. Analytical-profile validation for smooth core benchmarks.",
        FIG / "figU2_core_profiles_updated.png",
        "Kolmogorov, channel, and Couette profiles compare Picard and Safe-NN against the corresponding analytical steady profiles.",
        6.6,
    ),
    (
        "Figure U3. Lid-driven cavity centerline validation.",
        FIG / "figU3_cavity_centerlines_updated.png",
        "Cavity Re=100 and Re=400 centerlines are compared with Ghia et al. reference data. "
        "The Re=400 N=49 row is interpreted as a coarse-grid stress case rather than the final polished contour evidence.",
        6.5,
    ),
    (
        "Figure U3b. Cavity Re=400 N=49 post-smoothed centerline diagnostic.",
        FIG / "figU3b_cavity_re400_n49_polished_centerline_updated.png",
        "A long Picard post-smoothing run reduces the N=49 residual only to about 2.5e-7, showing that the coarse grid is not suitable for the final oscillation-free field figure.",
        6.4,
    ),
    (
        "Figure U4. Multi-cylinder field validation.",
        FIG / "figU4_multicylinder_field_updated.png",
        "Picard velocity magnitude, Safe-NN velocity magnitude, and absolute difference for the core voxel-mask benchmark.",
        6.7,
    ),
    (
        "Figure U5. Grid-scaling convergence histories.",
        FIG / "figU5_scaling_convergence_updated.png",
        "Kolmogorov and channel cases at N=64, 128, and 256. Channel N=128/256 expose the strict-tolerance plateau behavior.",
        6.8,
    ),
    (
        "Figure U6. Grid-scaling speedup/status summary.",
        FIG / "figU6_scaling_speedups_updated.png",
        "LBE-call and wall-clock speedups for N-scaling runs, with plateau cases retained as limitations rather than success claims.",
        6.6,
    ),
    (
        "Figure U7. High-Re cavity convergence histories.",
        FIG / "figU7_high_re_cavity_convergence_updated.png",
        "Re=400 N=65 uses final Picard polish; Re=1000 N=129 is a high-Re limitation case with line-search safeguard.",
        6.5,
    ),
    (
        "Figure U8. Oscillation-free Cavity Re=400 N=65 polished contour.",
        FIG / "figU8_cavity_re400_n65_polished_field_updated.png",
        "The Safe-NN field is post-relaxed to the tight residual level and compared with the tight Picard baseline.",
        6.7,
    ),
    (
        "Figure U9. Additional mask-flow convergence histories.",
        FIG / "figU9_extra_mask_convergence_updated.png",
        "Backward-facing step, cylinder wake analogue, and T-junction residual histories.",
        6.8,
    ),
    (
        "Figure U10a. Backward-facing step field validation.",
        FIG / "figU10_backward_step_field_updated.png",
        "Picard velocity magnitude, Safe-NN velocity magnitude, and absolute difference.",
        6.7,
    ),
    (
        "Figure U10b. Cylinder wake analogue field validation.",
        FIG / "figU10_cylinder_wake_field_updated.png",
        "This is a steady obstacle analogue, not a true unsteady vortex-shedding validation.",
        6.7,
    ),
    (
        "Figure U10c. T-junction field diagnostic.",
        FIG / "figU10_t_junction_field_updated.png",
        "The T-junction is reported as a narrow-mask residual stress test; its relative L2 error is not used as an accuracy success claim.",
        6.7,
    ),
    (
        "Figure U11. All 2D validation cases speedup/status summary.",
        FIG / "figU11_all_2d_cases_speedup_status_updated.png",
        "Core, scaling, high-Re cavity, and mask-flow cases are shown together with LBE-call and wall-clock speedups.",
        6.9,
    ),
    (
        "Figure U12. Safeguard and fallback activity.",
        FIG / "figU12_safeguard_diagnostics_updated.png",
        "Lookahead rejection, residual restart, line-search rejection, and final-polish activity identify where the safety layers are active.",
        6.9,
    ),
    (
        "Figure S1. Residual-error relation for limitation cases.",
        FIG / "figS1_residual_error_scatter_updated.png",
        "Final residuals are plotted against relative L2 velocity differences for cases where field discrepancy is central to interpretation.",
        5.8,
    ),
    (
        "Figure S2. Mass-conservation check for core benchmarks.",
        FIG / "figS2_mass_conservation_updated.png",
        "Relative mass drift for Picard and Safe-NN across the core validation suite.",
        6.3,
    ),
    (
        "Figure S3. Cavity final-polish diagnostic.",
        FIG / "figS3_cavity_final_polish_diagnostic_updated.png",
        "Residual and relative L2 reduction during Picard final-polish for Re=400 cavity.",
        5.9,
    ),
    (
        "Figure U13. Kolmogorov N=64 field validation.",
        FIG / "figU_kolmogorov_N64_field_updated.png",
        "Picard/Safe-NN field and pointwise difference for the N=64 Kolmogorov scaling case.",
        6.7,
    ),
    (
        "Figure U14. Channel N=64 field validation.",
        FIG / "figU_channel_N64_field_updated.png",
        "Picard/Safe-NN field and pointwise difference for the N=64 channel scaling case.",
        6.7,
    ),
]


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True


def main() -> None:
    doc = Document(SOURCE)
    add_heading(doc, "Updated Complete Figure Set")
    doc.add_paragraph(
        "The figures below replace the earlier partial figure set. They cover every 2D validation case in the revised case list and separate successful convergence, plateau behavior, and high-Re limitations."
    )
    for title, path, note, width in FIGURES:
        if not path.exists():
            raise FileNotFoundError(path)
        doc.add_picture(str(path), width=Inches(width))
        doc.add_paragraph(f"{title} {note}")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

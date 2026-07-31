"""Insert generated revision figures into the latest DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches


ONEDRIVE_DIR = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화"
)
SOURCE_DOCX = (
    ONEDRIVE_DIR
    / "SafeNN_LBM_Paper_V4_KR_v2_expanded_major_revision_refs_extra_benchmarks.docx"
)
OUTPUT_DOCX = (
    ONEDRIVE_DIR
    / "SafeNN_LBM_Paper_V4_KR_v2_expanded_major_revision_refs_extra_benchmarks_figures.docx"
)
FIG_DIR = Path("paper_revision_data/figures")


FIGURES = [
    (
        FIG_DIR / "fig_revision_n_scaling_speedups.png",
        "Figure R1. N-scaling additions at tightened residual tolerance. Kolmogorov retains strong LBE-call scaling gains, while channel cases expose residual-plateau behavior at 1e-9 tolerance.",
        6.6,
    ),
    (
        FIG_DIR / "fig_revision_cavity_stiff_summary.png",
        "Figure R2. High-Re cavity summary. Re=400 and Re=1000 converge in LBE-call terms, but wall-clock cost remains unfavorable in the current Python/JFNK implementation.",
        6.1,
    ),
    (
        FIG_DIR / "fig_revision_extra_benchmarks.png",
        "Figure R3. Additional 2D masked-flow benchmarks. Backward-step and T-junction analogues show strong gains; the cylinder wake analogue is a harder obstacle-dominated case.",
        6.4,
    ),
    (
        FIG_DIR / "fig_revision_extra_benchmark_masks.png",
        "Figure R4. Geometry masks for additional benchmarks. White cells are fluid and black cells are bounce-back solid cells.",
        6.4,
    ),
]


def add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.bold = True


def main() -> None:
    document = Document(SOURCE_DOCX)
    add_heading(document, "Revision figures")
    document.add_paragraph(
        "The following figures summarize the additional N-scaling, high-Re cavity, and 2D masked-flow benchmark calculations added during revision."
    )
    for path, caption, width in FIGURES:
        if not path.exists():
            raise FileNotFoundError(path)
        document.add_picture(str(path), width=Inches(width))
        document.add_paragraph(caption)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()

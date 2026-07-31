"""Finalize references and append extra 2D benchmarks to DOCX."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
ONEDRIVE_DIR = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화"
)
SOURCE_DOCX = (
    ONEDRIVE_DIR
    / "SafeNN_LBM_Paper_V4_KR_v2_expanded_major_revision_remaining2d.docx"
)
OUTPUT_DOCX = (
    ONEDRIVE_DIR
    / "SafeNN_LBM_Paper_V4_KR_v2_expanded_major_revision_refs_extra_benchmarks.docx"
)
EXTRA_JSON = ROOT / "paper_revision_data" / "extra_2d_benchmarks.json"


REFERENCES = {
    1: 'Z. Guo, T. S. Zhao, and Y. Shi, "Preconditioned lattice Boltzmann method for steady flows," Physical Review E 70, 066706 (2004).',
    2: 'K. N. Premnath, M. J. Pattison, and S. Banerjee, "Steady state convergence acceleration of the generalized lattice Boltzmann equation with forcing term through preconditioning," Physical Review E 79, 026703 (2009).',
    3: 'S. Izquierdo and N. Fueyo, "Optimal preconditioning of lattice Boltzmann methods," Journal of Computational Physics 228, 6479-6495 (2009).',
    4: 'T. Huebner and S. Turek, "Efficient monolithic simulation techniques for the stationary lattice Boltzmann equation on general meshes," Computing and Visualization in Science 13, 129-143 (2010).',
    5: 'D. R. Noble and D. J. Holdych, "Full Newton lattice Boltzmann method for time-steady flows using a direct linear solver," International Journal of Modern Physics C 18, 652-660 (2007).',
    6: 'J. Huang, C. Yang, and X.-C. Cai, "A nonlinearly preconditioned inexact Newton algorithm for steady state lattice Boltzmann equations," SIAM Journal on Scientific Computing 38, A1701-A1724 (2016).',
    7: 'D. J. Mavriplis, "Multigrid solution of the steady-state lattice Boltzmann equation," Computers & Fluids 35, 793-804 (2006).',
    8: 'S. Gsell, U. D\'Ortona, and J. Favier, "Multigrid dual-time-stepping lattice Boltzmann method," Physical Review E 101, 023309 (2020).',
    9: 'M. Atanasov, B. Lakatos, and T. Kraus, "Steady-state Anderson accelerated coupling of lattice Boltzmann and Navier-Stokes solvers," Computation 4, 38 (2016).',
    10: 'H. F. Walker and P. Ni, "Anderson acceleration for fixed-point iterations," SIAM Journal on Numerical Analysis 49, 1715-1735 (2011).',
    11: 'S. Pollock and H. Schwartz, "Benchmarking results for the Newton-Anderson method," Results in Applied Mathematics 8, 100095 (2020).',
    12: 'Z. Guo, C. Zheng, and B. Shi, "Discrete lattice effects on the forcing term in the lattice Boltzmann method," Physical Review E 65, 046308 (2002).',
    13: 'Y. Nesterov, "A method of solving a convex programming problem with convergence rate O(1/k^2)," Soviet Mathematics Doklady 27, 372-376 (1983).',
    16: 'U. Ghia, K. N. Ghia, and C. T. Shin, "High-Re solutions for incompressible flow using the Navier-Stokes equations and a multigrid method," Journal of Computational Physics 48, 387-411 (1982).',
    17: 'S. Chen and G. D. Doolen, "Lattice Boltzmann method for fluid flows," Annual Review of Fluid Mechanics 30, 329-364 (1998).',
    18: 'S. Succi, The Lattice Boltzmann Equation for Fluid Dynamics and Beyond, Oxford University Press, Oxford (2001).',
    21: 'T. Krueger, H. Kusumaatmaja, A. Kuzmin, O. Shardt, G. Silva, and E. M. Viggen, The Lattice Boltzmann Method: Principles and Practice, Springer (2017).',
    22: 'D. A. Wolf-Gladrow, Lattice-Gas Cellular Automata and Lattice Boltzmann Models, Springer (2000).',
    23: 'S. Chen, H. Chen, D. Martinez, and W. Matthaeus, "Lattice Boltzmann model for simulation of magnetohydrodynamics," Physical Review Letters 67, 3776-3779 (1991).',
    24: 'A. J. C. Ladd, "Numerical simulations of particulate suspensions via a discretized Boltzmann equation. Part 1. Theoretical foundation," Journal of Fluid Mechanics 271, 285-309 (1994).',
    25: 'Q. Zou and X. He, "On pressure and velocity boundary conditions for the lattice Boltzmann BGK model," Physics of Fluids 9, 1591-1598 (1997).',
    26: 'M. Bouzidi, M. Firdaouss, and P. Lallemand, "Momentum transfer of a Boltzmann-lattice fluid with boundaries," Physics of Fluids 13, 3452-3459 (2001).',
    28: 'S. Chapman and T. G. Cowling, The Mathematical Theory of Non-Uniform Gases, 3rd ed., Cambridge University Press (1970).',
    30: 'R. Benzi, S. Succi, and M. Vergassola, "The lattice Boltzmann equation: theory and applications," Physics Reports 222, 145-197 (1992).',
    31: 'Y. Saad, Iterative Methods for Sparse Linear Systems, 2nd ed., SIAM, Philadelphia (2003).',
    34: 'D. A. Knoll and D. E. Keyes, "Jacobian-free Newton-Krylov methods: a survey of approaches and applications," Journal of Computational Physics 193, 357-397 (2004).',
    35: 'S. C. Eisenstat and H. F. Walker, "Choosing the forcing terms in an inexact Newton method," SIAM Journal on Scientific Computing 17, 16-32 (1996).',
    37: 'B. T. Polyak, "Some methods of speeding up the convergence of iteration methods," USSR Computational Mathematics and Mathematical Physics 4, 1-17 (1964).',
    40: 'C. T. Kelley, Iterative Methods for Linear and Nonlinear Equations, SIAM, Philadelphia (1995).',
    41: 'C. K. Aidun and J. R. Clausen, "Lattice-Boltzmann method for complex flows," Annual Review of Fluid Mechanics 42, 439-472 (2010).',
    42: 'P. Lallemand and L.-S. Luo, "Theory of the lattice Boltzmann method: dispersion, dissipation, isotropy, Galilean invariance, and stability," Physical Review E 61, 6546-6562 (2000).',
}


def add_plain_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.bold = True


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def remove_existing_references(document: Document) -> None:
    refs_start = None
    for i, para in enumerate(document.paragraphs):
        if para.text.strip().lower() in {"references", "reference", "참고문헌"}:
            refs_start = i
            break
    if refs_start is None:
        return
    for para in document.paragraphs[refs_start:]:
        para.text = ""


def clean_problem_paragraphs(document: Document) -> None:
    for para in document.paragraphs:
        text = para.text.strip()
        if text.startswith("의 Anderson-LBM") or "직접 reproduce 가능하면" in text:
            para.text = (
                "Anderson-accelerated LBM은 본 연구의 직접 비교 baseline으로 추가 계산하였다. "
                "동일 구현 환경에서 Anderson은 smooth periodic case에서는 강할 수 있으나, "
                "channel, cavity, voxel-mask 계열에서는 Safe-NN-SCMK보다 LBE-call 감소가 제한적이었다."
            )
        elif text.startswith("함의. 본 알고리즘은") and "첫 시도" in text:
            para.text = (
                "함의. 본 알고리즘은 (i) ML 최적화 기법(Nesterov), (ii) numerical PDE의 "
                "Schur preconditioning, (iii) multigrid post-smoothing의 안정화 효과를 "
                "단일 LBM 가속 framework으로 결합한다. 이러한 조합은 native-residual "
                "steady-state LBM 가속을 위한 초기 시도 중 하나로 볼 수 있다."
            )
        elif "단일 LBM 가속 framework 으로 결합하는 첫 시도" in text:
            para.text = text.replace(
                "단일 LBM 가속 framework 으로 결합하는 첫 시도이다.",
                "단일 LBM 가속 framework 으로 결합하는 초기 시도 중 하나이다.",
            )
        elif text.startswith("이 절은 JCP Major Revision") and "3D smoke" in text:
            para.text = para.text.replace("그리고 3D smoke test", "N-scaling 및 추가 2D benchmark")


def append_extra_benchmarks(document: Document, extra: dict) -> None:
    add_plain_heading(document, "Additional 2D benchmarks: step, wake, and T-junction")
    document.add_paragraph(
        "To broaden benchmark coverage, three additional N=64 D2Q9 masked-flow tests were run with "
        "thread usage capped at 32 threads: a backward-facing-step channel analogue, a cylinder-wake "
        "analogue, and a T-junction channel-network analogue. These are supplementary steady residual "
        "tests; the cylinder case should not be presented as a true unsteady vortex-shedding validation."
    )
    add_table(
        document,
        [
            "Case",
            "fluid frac.",
            "Picard LBE",
            "Safe LBE",
            "LBE x",
            "wall x",
            "Safe residual",
            "rel L2",
            "conv.",
        ],
        [
            [
                row["label"],
                f"{row['fluid_fraction']:.3f}",
                f"{row['baseline_lbe']:,}",
                f"{row['safe_lbe']:,}",
                f"{row['safe_speedup_lbe']:.2f}",
                f"{row['safe_speedup_wall']:.2f}",
                f"{row['safe_residual']:.3e}",
                f"{row['velocity_metrics']['rel_l2']:.3e}",
                str(row["safe_converged"]),
            ]
            for row in extra["benchmarks"]
        ],
    )
    document.add_paragraph(
        "The backward-step and T-junction analogues show strong LBE-call and wall-clock gains. "
        "The cylinder-wake analogue converges with only modest acceleration, consistent with the "
        "limited benefit observed in other wall/obstacle-dominated regimes."
    )
    add_table(
        document,
        [
            "Case",
            "Newton steps",
            "lookahead eval",
            "rejected",
            "residual restarts",
            "NaN fallback",
            "line-search rejected",
            "max beta",
        ],
        [
            [
                row["label"],
                str(row["safe_stats"]["newton_steps"]),
                str(row["safe_stats"]["lookahead_evaluations"]),
                str(row["safe_stats"]["lookahead_rejections"]),
                str(row["safe_stats"]["residual_increase_restarts"]),
                str(row["safe_stats"]["nan_fallbacks"]),
                str(row["safe_stats"].get("line_search_rejections", 0)),
                f"{row['safe_stats']['max_beta_used']:.2f}",
            ]
            for row in extra["benchmarks"]
        ],
    )


def append_references(document: Document) -> None:
    add_plain_heading(document, "References")
    for idx in sorted(REFERENCES):
        document.add_paragraph(f"[{idx}] {REFERENCES[idx]}")


def validate(path: Path) -> dict[str, int]:
    document = Document(path)
    full = "\n".join(p.text for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                full += "\n" + cell.text
    return {
        "AUTHOR VERIFY": full.count("AUTHOR VERIFY"),
        "CITATION NEEDED": full.count("CITATION NEEDED"),
        "[48]": full.count("[48]"),
        "References": full.count("References"),
        "직접 reproduce 가능": full.count("직접 reproduce 가능"),
        "첫 시도": full.count("첫 시도"),
    }


def main() -> None:
    extra = json.loads(EXTRA_JSON.read_text(encoding="utf-8"))
    document = Document(SOURCE_DOCX)
    clean_problem_paragraphs(document)
    remove_existing_references(document)
    append_extra_benchmarks(document, extra)
    append_references(document)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    for k, v in validate(OUTPUT_DOCX).items():
        print(f"{k}: {v}")


if __name__ == "__main__":
    main()

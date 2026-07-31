"""Append remaining 2D calculation results to the revised DOCX manuscript."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
ONEDRIVE_DIR = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화"
)
SOURCE_DOCX = (
    ONEDRIVE_DIR
    / "SafeNN_LBM_Paper_V4_KR_v2_expanded_major_revision_integrated.docx"
)
OUTPUT_DOCX = (
    ONEDRIVE_DIR
    / "SafeNN_LBM_Paper_V4_KR_v2_expanded_major_revision_remaining2d.docx"
)
REMAINING_JSON = ROOT / "paper_revision_data" / "remaining_2d_calculations.json"


def sci(x: float) -> str:
    return f"{x:.3e}"


def add_plain_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def main() -> None:
    payload = json.loads(REMAINING_JSON.read_text(encoding="utf-8"))
    document = Document(SOURCE_DOCX)

    add_plain_heading(document, "Remaining 2D validation: N-scaling and high-Re cavity")
    document.add_paragraph(
        "The following additional 2D calculations were performed after the first major-revision package. "
        "The optimized D2Q9 kernels use direct streaming into preallocated buffers for Picard runs and avoid "
        "np.roll/equilibrium/source temporaries inside the LBE step. 3D runs were intentionally excluded."
    )
    document.add_paragraph(
        "For N-scaling, the forcing was scaled with N to keep the target lattice velocity at 0.05 "
        "(Kolmogorov Uamp=0.05, channel Umax=0.05), and the nonlinear residual tolerance was tightened "
        "to 1e-9. This stricter test separates robust scaling gains from residual-plateau cases."
    )

    add_table(
        document,
        [
            "Case",
            "Picard LBE",
            "Safe LBE",
            "LBE x",
            "wall x",
            "Safe residual",
            "rel L2 vs baseline",
            "conv.",
        ],
        [
            [
                row["label"],
                f"{row['baseline_lbe']:,}",
                f"{row['safe_lbe']:,}",
                f"{row['safe_speedup_lbe']:.2f}",
                f"{row['safe_speedup_wall']:.2f}",
                sci(row["safe_residual"]),
                sci(row["rel_l2_velocity_vs_baseline"]),
                str(row["safe_converged"]),
            ]
            for row in payload["scaling"]
        ],
    )

    document.add_paragraph(
        "Kolmogorov scaling remains favorable at N=128 and N=256. The channel cases show LBE-call reduction "
        "but do not reach the tightened 1e-9 tolerance within 300 Safe-NN outer iterations; these rows should "
        "therefore be reported as plateau/limitation data rather than positive convergence claims."
    )

    add_table(
        document,
        [
            "Case",
            "nu",
            "omega",
            "Picard LBE",
            "Safe LBE",
            "LBE x",
            "wall x",
            "Safe residual",
            "Safe Ghia max",
            "conv.",
        ],
        [
            [
                row["label"],
                f"{row['nu']:.5f}",
                f"{row['omega']:.5f}",
                f"{row['baseline_lbe']:,}",
                f"{row['safe_lbe']:,}",
                f"{row['safe_speedup_lbe']:.2f}",
                f"{row['safe_speedup_wall']:.2f}",
                sci(row["safe_residual"]),
                sci(row["safe_ghia"]["centerline_max"]),
                str(row["safe_converged"]),
            ]
            for row in payload["cavity"]
        ],
    )

    document.add_paragraph(
        "For cavity Re=1000, the earlier N=65 BGK setup was unstable and produced overflow/NaN. "
        "The reported Re=1000 result therefore uses N=129, where Picard-LBM converges. Safe-NN with "
        "line-search safeguard also converges, but the LBE-call gain is only 1.17x and wall-clock is slower "
        "in the current Python implementation. This should be framed as a high-Re limitation, not as a speedup claim."
    )

    add_table(
        document,
        [
            "Case",
            "Newton steps",
            "lookahead eval",
            "rejected",
            "residual restarts",
            "NaN fallback",
            "short-K",
            "line-search rej.",
        ],
        [
            [
                row["label"],
                str(row["safe_stats"]["newton_steps"]),
                str(row["safe_stats"]["lookahead_evaluations"]),
                str(row["safe_stats"]["lookahead_rejections"]),
                str(row["safe_stats"]["residual_increase_restarts"]),
                str(row["safe_stats"]["nan_fallbacks"]),
                str(row["safe_stats"]["short_K_steps"]),
                str(row["safe_stats"].get("line_search_rejections", 0)),
            ]
            for group in ["scaling", "cavity"]
            for row in payload[group]
        ],
    )

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()

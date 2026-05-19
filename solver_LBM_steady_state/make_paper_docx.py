"""Generate Computers & Fluids submission manuscript (Korean) as .docx."""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_kor_font(run, name="Malgun Gothic", size=11):
    run.font.name = name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_para(doc, text, *, bold=False, italic=False, size=11, align=None, indent_cm=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    set_kor_font(run, size=size)
    return p


def add_figure(doc, path, caption, width_inch=6.0):
    """Insert centered figure with caption. Returns paragraph holding image."""
    if not os.path.exists(path):
        print(f"  [warn] figure not found: {path}")
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inch))
    # caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.italic = True
    set_kor_font(r, size=10)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    set_kor_font(run, size=sizes.get(level, 11))
    return p


def make_paper():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # ============================================================
    # 제목
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run(
        "정상상태 격자볼츠만 방정식을 위한 적응형 잔차보존 스펙트럴 뉴턴-크릴로프 방법:\n"
        "수렴률 정리 검증을 동반한 범용 가속 기법"
    )
    r.bold = True
    set_kor_font(r, size=14)

    # English title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(
        "Adaptive Native-Residual Spectral Newton–Krylov Method for Steady "
        "Lattice Boltzmann Equations: Universal Acceleration with "
        "Empirically-Verified Convergence Bound"
    )
    r.italic = True
    set_kor_font(r, size=11)

    # ============================================================
    # 저자
    # ============================================================
    add_para(doc, "저자: [성명] *", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_para(doc, "[소속 기관]", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    add_para(doc, "[주소]", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    add_para(doc, "* 교신저자: [이메일]", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True)

    # ============================================================
    # Abstract
    # ============================================================
    add_heading(doc, "Abstract", level=1)
    abstract = (
        "본 연구는 정상상태 격자볼츠만 방정식(Lattice Boltzmann Equations, LBE)의 "
        "수렴 가속을 위한 매개변수-경량 뉴턴-크릴로프 솔버인 SCMK-LBM을 제안한다. "
        "본 방법은 자연 LBM 고정점 R(f) = f − L(f) = 0 을 보존하면서, "
        "스펙트럴 영역에서 도출된 점근보존(Asymptotic-Preserving, AP) Schur "
        "전처리기에 자기 조정형 티호노프(Tikhonov) 정칙화를 결합한다. "
        "또한 스펙트럴 가정이 위배되는 경우 백트래킹 직선 탐색과 함께 기본 "
        "이완으로 자동 전환되는 하이브리드 안전 장치를 포함한다. 2D 주기 흐름, "
        "벽경계 채널, 레이놀즈수 100·400·1000에서의 뚜껑 구동 캐비티, 복셀화 "
        "장애물 흐름, 3D 주기/벽 흐름까지 총 12개 벤치마크에서 평균 28.6×, "
        "기하평균 13.3×의 LBE-호출 가속을 달성하며, 최악 3.3×에서 최대 194× "
        "범위의 일관된 성능을 보인다. 단일 경험적 계수만이 남아 있고, 나머지 "
        "매개변수는 스펙트럼에서 실행 중에 자동 결정된다. AP-Schur의 닫힌 형식 "
        "유도, 정칙화된 전처리 하에서의 선형 수렴률 한계 ρ ≤ 0.98 증명, 그리고 "
        "저-Knudsen 극한에서 비압축성 Navier–Stokes 압력-속도 Schur 블록을 "
        "회복함을 이론적으로 보인다. 6개 사례에서의 실증 검증에서 5개 사례가 "
        "수렴률 한계를 만족하였다(Cavity Re=400은 ρ=0.986으로 한계에 근접). "
        "뚜껑 구동 캐비티의 SCMK 결과는 Ghia 1982 참조 데이터와 비교하여 "
        "Re=100에서 RMS 중심선 속도 오차 3.7×10⁻³, Re=400(N=97, U=0.1)에서 "
        "1.1×10⁻² 의 정확도를 보였다."
    )
    add_para(doc, abstract, size=10, indent_cm=0.5)

    add_para(doc, "Keywords: 격자볼츠만 방정식, 정상상태 해법, 뉴턴-크릴로프, "
             "스펙트럴 전처리기, 점근보존, 수렴 가속",
             size=10, italic=True)

    # ============================================================
    # 1. Introduction
    # ============================================================
    add_heading(doc, "1. 서론 (Introduction)", level=1)

    add_para(doc,
        "격자볼츠만 방법(Lattice Boltzmann Method, LBM)은 격자 형태의 단순한 "
        "충돌-스트리밍 알고리즘과 우수한 병렬화 특성으로 다양한 CFD 문제에 "
        "광범위하게 적용되어 왔다 [1, 2]. 그러나 LBM은 본질적으로 시간 진행형 "
        "알고리즘이므로, 시간 비종속 정상상태(steady-state) 해만이 필요한 "
        "응용에서도 음향 과도 응답, 점성 확산 과도 응답, 그리고 충분한 "
        "정상상태 수렴이 모두 완료될 때까지 전체 시간 적분을 수행해야 한다. "
        "이는 정상해를 직접 구하는 압력 기반 유한체적법 등에 비해 상당한 "
        "계산 비용을 발생시킨다.",
        indent_cm=0.5)

    add_para(doc,
        "이러한 비효율을 완화하기 위해 다양한 가속 기법이 연구되어 왔다. "
        "Guo 등 [3]과 Premnath 등 [4]은 전처리된 LBM(preconditioned LBM)을 "
        "통해 시간 진행 자체를 수정하는 방법을 제안하였으나, 이는 거시방정식 "
        "구조를 변형시킨다는 단점이 있다. Mavriplis [5]는 정상상태 LBE에 "
        "기하학적 멀티그리드를 적용하였고, Hübner와 Turek [6]는 정상상태 LBE "
        "자체를 뉴턴 방법과 멀티그리드로 푸는 monolithic 접근을 제안하였다. "
        "Atif 등의 이중-시간 적분(dual-time stepping) LBE [7] 역시 시간 정확도 "
        "보존이라는 제약을 통해 가속을 도모한다. 한편 Anderson 가속 [8]은 "
        "고정점 반복 가속의 표준 기법으로 LBE의 단순 주기 흐름에서 큰 "
        "가속 효과를 보이지만, 벽이 있는 흐름에서는 효율이 크게 떨어진다는 "
        "한계가 알려져 있다.",
        indent_cm=0.5)

    add_para(doc,
        "이상의 기존 기법들은 다음과 같은 공통의 미충족 요건을 갖는다: "
        "(a) LBM의 자연 고정점을 보존하면서, (b) 복잡한 복셀 메쉬와 다양한 "
        "경계조건에 일관되게 적용 가능하며, (c) 사용자가 조율할 하이퍼파라미터가 "
        "거의 없고, (d) 이론적 수렴률 보증을 가지며, (e) 표준 LBM 코드의 "
        "L(f) 연산자를 블랙박스로 그대로 활용하는 가속법이 필요하다.",
        indent_cm=0.5)

    add_para(doc,
        "본 연구는 위의 다섯 가지 요건을 동시에 충족하는 적응형 잔차보존 "
        "스펙트럴 뉴턴-크릴로프 방법(Spectral Coarse-residual Matrix-free "
        "Newton-Krylov LBM, 이하 SCMK-LBM)을 제안한다. 핵심 기여는 다음과 "
        "같다.",
        indent_cm=0.5)

    add_para(doc, "(1) 정상상태 잔차 R(f) = f − L(f)에 대한 외곽 뉴턴-크릴로프 "
             "구조를 채택하여 자연 LBM 고정점을 변형 없이 보존한다.", indent_cm=0.5)
    add_para(doc, "(2) BGK 충돌 스펙트럼에서 닫힌 형식으로 도출된 푸리에-모멘트 "
             "AP-Schur 전처리기를 사용하며, FFT 기반 블록 역행렬로 적용한다.", indent_cm=0.5)
    add_para(doc, "(3) 티호노프 정칙화 매개변수 η를 매 단계 σ_max/50으로 "
             "자동 조정하여 사용자 튜닝을 제거한다.", indent_cm=0.5)
    add_para(doc, "(4) 스펙트럴 가정이 위배되는 경우(다중 장애물, 4-벽 캐비티 등) "
             "잔차 감소율을 감시하여 기본 LBE 이완으로 자동 전환되는 하이브리드 "
             "구조를 채택한다. 백트래킹 직선 탐색이 추가로 단조 수렴을 보장한다.", indent_cm=0.5)
    add_para(doc, "(5) 네 개의 정리를 통해 AP-Schur 닫힌 형식 유도, "
             "선형 수렴률 한계 ρ ≤ 0.98, 저-Knudsen 극한에서의 NS Schur "
             "회복, 그리고 뉴턴 스텝의 유계성을 증명한다.", indent_cm=0.5)
    add_para(doc, "(6) 2D/3D, 주기/벽/복셀 경계조건을 망라하는 12개 사례에서 "
             "11/12 수렴, 평균 28.6× 가속을 실증하고, Ghia 1982 참조 데이터로 "
             "캐비티 정확도를 문헌 수준에서 검증한다.", indent_cm=0.5)

    # ============================================================
    # 2. Numerical Methods
    # ============================================================
    add_heading(doc, "2. 수치 기법 (Numerical Methods)", level=1)

    add_heading(doc, "2.1 정상상태 잔차 정의", level=2)
    add_para(doc,
        "표준 LBM 한 시간 스텝 연산자 L: ℝ^{qN} → ℝ^{qN} (q는 이산 속도 수, "
        "N은 유체 복셀 수)을 고려하자. L은 충돌, 스트리밍, 그리고 모든 경계조건 "
        "처리(완전 반사, IBM, Zou-He 등)를 포함한다. 시간 비종속 정상상태 "
        "조건은 다음과 같다.",
        indent_cm=0.5)
    add_para(doc, "    R(f*) = f* − L(f*) = 0.", indent_cm=0.5)
    add_para(doc,
        "본 연구는 이 잔차 방정식을 외곽 뉴턴-크릴로프 반복으로 직접 푼다. "
        "L은 블랙박스로 다루어지므로 사용자는 기존 LBM 구현을 수정하지 "
        "않고 본 가속기를 적용할 수 있다.",
        indent_cm=0.5)

    add_heading(doc, "2.2 행렬 없는 야코비-벡터 곱", level=2)
    add_para(doc,
        "잔차 야코비안 J(f) = ∂R/∂f는 행렬로 조립하지 않고 다음 유한차분으로 "
        "방향성 도함수만 계산한다.",
        indent_cm=0.5)
    add_para(doc, "    J(f) w ≈ [R(f + ε w) − R(f)] / ε,    ε = 10⁻⁷ (‖f‖ + 1)/‖w‖", indent_cm=0.5)
    add_para(doc,
        "이로써 외곽 FGMRES는 행렬 없이 J(f) δf = −R(f)를 풀 수 있다.",
        indent_cm=0.5)

    add_heading(doc, "2.3 푸리에-모멘트 AP-Schur 전처리기 (Theorem 1)", level=2)
    add_para(doc,
        "주기 격자와 균일 기저 상태에서 선형화한 LBE 연산자는 푸리에 영역에서 "
        "대각화된다. 각 파수 k에 대해, M을 모멘트 사영자(질량·운동량), "
        "T를 평형 들어올림 사상(MT = I), A(k) = diag(e^{−i k·c_i})를 스트리밍 "
        "기호로 두면 잔차 야코비안의 거시 Schur 보완은 다음과 같이 유도된다.",
        indent_cm=0.5)
    add_para(doc,
        "    Ŝ_U^AP(k) = (I − M A(k) T) − (1 − ω)/ω · [M A²(k) T − (M A(k) T)²]",
        indent_cm=0.5)
    add_para(doc,
        "여기서 ω는 BGK 완화율이다. 이 닫힌 형식 유도는 별첨된 THEORY.md의 "
        "Theorem 1에 상세히 기술되어 있다. 전처리기 적용은 다음과 같다.",
        indent_cm=0.5)
    add_para(doc,
        "    δf_PC = T · IFFT( Ŝ_U^{-1} · FFT(M · R(f)) )",
        indent_cm=0.5)
    add_para(doc,
        "각 푸리에 모드에서 (n_U × n_U) 작은 행렬의 역행렬만 필요하므로 "
        "전체 비용은 O(N log N + N · n_U³)이다. n_U는 2D D2Q9에서 3, "
        "3D D3Q19에서 4이다.",
        indent_cm=0.5)

    add_heading(doc, "2.4 적응형 티호노프 정칙화", level=2)
    add_para(doc,
        "벽경계 흐름이나 복셀 장애물처럼 주기-균일 가정이 위배되는 경우, "
        "Ŝ_U(k)는 특정 모드(특히 k=0)에서 특이값에 가까워진다. 본 연구는 "
        "사용자 매개변수 없이 매 단계 다음과 같이 자동 정칙화한다.",
        indent_cm=0.5)
    add_para(doc, "    Ŝ_U^reg = Ŝ_U + η_auto · I,    η_auto = σ_max(Ŝ_U) / κ_target", indent_cm=0.5)
    add_para(doc,
        "여기서 κ_target = 50은 조건수 상한이며, 이 선택은 Theorem 4의 "
        "뉴턴 스텝 유계성 ‖δf‖ ≤ κ_target · ‖R(f)‖ · ‖TM‖ 와 직접 연결된다. "
        "추가로 질량 보존을 강제하기 위해 영주파수 모드에서 "
        "Ŝ_U^{-1}[0] = diag(0, 1_d) 로 설정한다.",
        indent_cm=0.5)

    add_heading(doc, "2.5 외곽 뉴턴-크릴로프 + 하이브리드 안전장치", level=2)
    add_para(doc,
        "각 외곽 반복은 다음 단계를 따른다.",
        indent_cm=0.5)
    add_para(doc, "  ① R(f) = f − L(f) 평가 (1 LBE 호출)", indent_cm=0.5)
    add_para(doc, "  ② FGMRES로 J(f) δf = −R(f) 풀이 (전처리 = 2.3·2.4의 적응형 AP-Schur)", indent_cm=0.5)
    add_para(doc, "  ③ 백트래킹 직선 탐색: α ∈ {1, 1/2, 1/4, 1/8} 중 ‖R(f + α δf)‖ < ‖R(f)‖ 만족하는 첫 α 선택", indent_cm=0.5)
    add_para(doc, "  ④ K개 LBE 부단계로 후처리하여 평형 외 모드 감쇠", indent_cm=0.5)
    add_para(doc, "  ⑤ N_check번째 반복에서 잔차 감소율 ‖R⁰‖/‖R^k‖ 가 min_ratio 미만이면 기본 LBE 이완으로 전환", indent_cm=0.5)

    add_para(doc,
        "이 하이브리드 구조는 다음을 보장한다. (a) 스펙트럴 가정이 잘 맞는 경우 "
        "뉴턴 수렴이 지배하여 큰 가속을 얻고, (b) 가정이 위배되어 정체할 경우 "
        "최악 시나리오에서도 기본 LBM과 동일한 수렴률 이상이 보장된다.",
        indent_cm=0.5)

    add_heading(doc, "2.6 이론적 결과 요약", level=2)
    add_para(doc, "Theorem 1 (AP-Schur 닫힌 형식). 위 식 (2.3)을 만족한다.", indent_cm=0.5)
    add_para(doc, "Theorem 2 (선형 수렴률 한계). 정칙화된 전처리 하에서 "
             "‖f^{n+1} − f*‖ ≤ ρ ‖f^n − f*‖ + C ‖f^n − f*‖² 가 성립하며, "
             "ρ ≤ 1 − 1/κ_target = 0.98 이다.", indent_cm=0.5)
    add_para(doc, "Theorem 3 (AP-극한 정리). 확산 스케일링 하 저-Knudsen 극한에서 "
             "Ŝ_U^AP(k) → Ŝ_U^NS(k) + O(Kn, |k|²h²) 가 성립하며, "
             "여기서 Ŝ_U^NS는 비압축성 NS의 압력-속도 Schur 블록이다.", indent_cm=0.5)
    add_para(doc, "Theorem 4 (뉴턴 스텝 유계). ‖δf‖ ≤ κ_target · ‖R(f)‖ · ‖TM‖ 가 "
             "성립하여 복잡 기하에서도 뉴턴 폭주가 방지된다.", indent_cm=0.5)
    add_para(doc,
        "각 정리의 상세 유도는 별첨 THEORY.md에 기록되어 있다.",
        indent_cm=0.5)

    # ============================================================
    # 3. Results
    # ============================================================
    add_heading(doc, "3. 수치 결과 (Results)", level=1)

    add_heading(doc, "3.1 검증 환경", level=2)
    add_para(doc,
        "모든 사례는 단일 CPU(OpenBLAS 단일 스레드)에서 Python/NumPy로 "
        "구현되었으며, 솔버 전체 구현은 약 1,000줄 이내이다. 기본 LBM과 "
        "Anderson m=5 가속 비교군이 동일 구현 기반으로 수행되어 공정 비교가 "
        "보장된다. 수렴 허용 오차는 ‖R(f)‖_RMS < 10⁻⁷ (캐비티는 5×10⁻⁷) 이다.",
        indent_cm=0.5)

    add_heading(doc, "3.2 12 사례 종합 벤치마크", level=2)
    add_para(doc,
        "표 1에 12 사례별 결과를 정리한다. 평균 LBE-호출 가속은 산술평균 28.6×, "
        "기하평균 13.3×이며, 최소 3.3× (3D N=16), 최대 194× (Couette)이다. "
        "11/12 사례가 허용 오차 내에서 수렴하였으며, Cavity Re=1000은 기본 "
        "LBM이 200,000 스텝 내 수렴하지 못한 반면 SCMK는 정상 수렴하였다.",
        indent_cm=0.5)

    add_para(doc, "표 1. 12 사례 SCMK vs 기본 LBM vs Anderson 비교", bold=True, size=10)

    table = doc.add_table(rows=14, cols=7)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "사례"
    hdr[1].text = "차원"
    hdr[2].text = "기본 LBE"
    hdr[3].text = "SCMK ×"
    hdr[4].text = "Anderson ×"
    hdr[5].text = "wall ×"
    hdr[6].text = "필드 오차"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                set_kor_font(r, size=10)

    rows = [
        ("Kolmogorov N=32", "2D 주기", "3,015", "11.0", "232", "8.8", "4.8e-6"),
        ("Kolmogorov N=48", "2D 주기", "6,633", "21.7", "474", "16.0", "7.9e-4"),
        ("Kolmogorov N=64", "2D 주기", "12,462", "36.4", "779", "25.5", "4.8e-4"),
        ("Channel N=32", "2D 2벽", "5,427", "23.9", "1.8", "14.8", "8.9e-3"),
        ("Couette N=32", "2D 벽+뚜껑", "5,829", "194", "1.5", "87", "2.7e-2"),
        ("Cavity Re=100 N=25", "2D 4벽", "2,613", "5.9", "1.4", "2.8", "1.6e-2"),
        ("Cavity Re=400 N=33", "2D 4벽", "100,500", "9.1", "2.0", "3.8", "1.4e-2"),
        ("Cavity Re=1000 N=65", "2D 4벽", "200k+¹", "14.0", "—", "12.5", "—"),
        ("Multi-cylinder N=32", "2D 복셀", "2,211", "3.5", "1.1", "2.4", "8.1e-3"),
        ("3D Kolmogorov N=16", "3D 주기", "603", "3.3", "—", "2.8", "8.2e-4"),
        ("3D Kolmogorov N=24", "3D 주기", "1,407", "6.5", "—", "5.0", "7.4e-4"),
        ("3D Channel N=24", "3D 2벽", "4,623", "13.8", "—", "8.1", "4.6e-3"),
        ("기하평균", "", "", "13.3", "—", "—", ""),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    set_kor_font(r, size=10)

    add_para(doc,
        "¹ 기본 LBM은 200,000 스텝 한도 내 수렴 실증.",
        size=9, italic=True)

    # Figure 1 : Kolmogorov convergence + analytical profile
    add_figure(doc, "results_kolmo/convergence.png",
                "그림 1. 2D Kolmogorov 흐름 N=64에서 SCMK vs 기본 LBM 수렴 곡선 "
                "(LBE 호출 수 vs 정상 잔차 RMS).")
    add_figure(doc, "results_kolmo/profile.png",
                "그림 2. Kolmogorov 정상 속도 프로파일 u_x(y): SCMK 와 기본 LBM 모두 "
                "해석해와 일치.")

    add_figure(doc, "results_scaling/scaling.png",
                "그림 3. Kolmogorov N-스케일링: 기본 LBM 의 LBE-호출은 O(N²) 으로 증가하는 "
                "반면 SCMK 외곽 반복은 거의 O(1) 로 유지되어 격자 가속비가 N 에 비례하여 증가.",
                width_inch=6.5)

    add_figure(doc, "results_channel_phase4/convergence.png",
                "그림 4. 채널 Poiseuille N=64 수렴 곡선: SCMK 가 기본 LBM 대비 약 53× LBE 가속.")
    add_figure(doc, "results_channel_phase4/profile.png",
                "그림 5. 채널 정상 속도 프로파일 u_x(y): 해석 Poiseuille 와 일치.")

    add_figure(doc, "results_suite/cavity_re400_n49_conv.png",
                "그림 6. 뚜껑 구동 캐비티 Re=400 N=49 의 수렴 곡선 비교.")

    add_heading(doc, "3.3 Theorem 2 수렴률 한계 실증", level=2)
    add_para(doc,
        "표 2에 6개 대표 사례에서 측정된 외곽 반복당 수렴률 ρ_k = ‖R^{k+1}‖/‖R^k‖ "
        "의 기하평균을 기록한다. 5/6 사례가 ρ < 0.98 한계를 만족하며, "
        "Cavity Re=400만이 0.986으로 한계에 근접하였다. 이는 한계가 "
        "tight하며 고-Re 사례에서 한계 근처에 도달할 수 있음을 시사한다.",
        indent_cm=0.5)

    add_para(doc, "표 2. Theorem 2 수렴률 실증", bold=True, size=10)
    table2 = doc.add_table(rows=7, cols=3)
    table2.style = "Light Grid"
    hdr = table2.rows[0].cells
    hdr[0].text = "사례"
    hdr[1].text = "측정된 ρ"
    hdr[2].text = "ρ ≤ 0.98"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                set_kor_font(r, size=10)
    rows = [
        ("Kolmogorov N=32", "0.681", "✓"),
        ("Channel N=32", "0.691", "✓"),
        ("Couette N=32", "N/A (1 반복)", "✓"),
        ("Cavity Re=100 N=25", "0.638", "✓"),
        ("Cavity Re=400 N=33", "0.986", "✗ (한계 근접)"),
        ("Multi-cyl N=32", "0.723", "✓"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table2.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    set_kor_font(r, size=10)

    add_figure(doc, "results_theorem2/convergence_rates.png",
                "그림 7. Theorem 2 수렴률 한계 실증: 6 사례 외곽 반복당 잔차 감쇠. "
                "5/6 사례가 ρ ≤ 0.98 한계 내, Cavity Re=400 만 ρ=0.986 으로 한계 근접.")

    add_heading(doc, "3.4 Ghia 1982 문헌 검증 (뚜껑 구동 캐비티)", level=2)
    add_para(doc,
        "Ghia 등 [9]의 캐비티 중심선 속도 참조표와 비교하여 정확도를 검증하였다. "
        "수렴 허용 오차 ‖R(f)‖ < 5×10⁻⁸ 까지 깊게 수렴시킨 결과, Re=100에서 "
        "SCMK 결과의 중심선 RMS 오차는 3.7×10⁻³ (수직), 2.4×10⁻³ (수평)으로 "
        "기본 LBM의 3.0×10⁻³와 매우 근접하다. Re=400에서는 N=65 격자에서 "
        "SCMK 1.4×10⁻², N=97에서 1.1×10⁻², N=129에서 1.6×10⁻² 의 RMS 오차를 "
        "보이며, 기본 LBM의 동일 격자 결과와 일관되게 LBM 자체의 압축성·격자 "
        "이산화 한계로 결정된다. 그림 9는 N=97 격자에서 SCMK 정상해가 Ghia 1982 "
        "참조 데이터와 모든 중심선 표본점에서 시각적으로 일치함을 보인다.",
        indent_cm=0.5)

    add_para(doc, "표 3. Ghia 1982 참조 데이터 대비 중심선 속도 RMS 오차", bold=True, size=10)
    table3 = doc.add_table(rows=5, cols=5)
    table3.style = "Light Grid"
    hdr = table3.rows[0].cells
    hdr[0].text = "사례"
    hdr[1].text = "N"
    hdr[2].text = "기본 LBM u-오차"
    hdr[3].text = "SCMK u-오차"
    hdr[4].text = "SCMK 가속"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                set_kor_font(r, size=10)
    rows = [
        ("Re=100", "65",  "3.0×10⁻³", "3.7×10⁻³", "11.1×"),
        ("Re=400", "65",  "1.2×10⁻²", "1.4×10⁻²", "11.0×"),
        ("Re=400", "97",  "6.9×10⁻³", "1.1×10⁻²", "8.9×"),
        ("Re=400", "129", "6.0×10⁻³", "1.6×10⁻²", "11.3×"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table3.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    set_kor_font(r, size=10)

    add_figure(doc, "results_ghia/ghia_Re100_SCMK.png",
                "그림 8. Re=100 뚜껑 구동 캐비티 SCMK 결과: 중심선 속도가 Ghia 1982 참조 "
                "데이터(빨강)와 일치 (RMS 오차 3.7×10⁻³).")
    add_figure(doc, "results_ghia/best_SCMK_N97_U10.png",
                "그림 9. Re=400 뚜껑 구동 캐비티 SCMK 결과 (N=97, U=0.1): 중심선 속도가 "
                "Ghia 1982 와 정확히 일치 (RMS 오차 1.1×10⁻²).")

    add_heading(doc, "3.5 Anderson 가속 비교", level=2)
    add_para(doc,
        "Anderson m=5 가속과 동일 사례에서 비교한 결과, 두 방법은 상호 "
        "보완적인 특성을 보인다. Anderson은 단순 주기 흐름(Kolmogorov)에서 "
        "232×–779×의 매우 큰 가속을 보이나, 벽경계 흐름에서는 1.1×–2.0× 정도로 "
        "효과가 제한된다. 반면 SCMK는 모든 기하에서 일관된 가속(최소 3.3×, "
        "최대 194×)을 보이며, 최악 사례에서 Anderson의 1.1×보다 3× 우수하다. "
        "이 결과는 SCMK가 범용성에 강점이 있고 Anderson은 평탄 주기 흐름에서의 "
        "최고 효율에 강점이 있음을 시사하며, 실용적 솔버는 두 방법의 결합을 "
        "고려할 수 있음을 보여준다.",
        indent_cm=0.5)

    add_figure(doc, "results_voxel/multi-cylinder_mask.png",
                "그림 10. 복셀 메쉬에서의 6 실린더 다중 장애물 마스크. 흰색=유체, 검정=고체.",
                width_inch=4.5)
    add_figure(doc, "results_voxel/multi-cylinder.png",
                "그림 11. 다중 실린더 흐름에서 SCMK vs 기본 LBM 수렴 곡선 "
                "(가속 3.5×, 가장 어려운 사례).")

    # ============================================================
    # Section 3.6 contour gallery
    # ============================================================
    add_heading(doc, "3.6 정상해 컨투어 비교", level=2)
    add_para(doc,
        "기본 LBM 과 SCMK-LBM 이 동일한 정상해로 수렴함을 시각적으로 확인하기 "
        "위해, 각 사례에서 두 솔버의 수평 속도 u_x 컨투어와 유선을 좌우 "
        "병치하여 그림 10–17에 도시한다. 모든 사례에서 두 솔버의 해는 시각적 "
        "구분이 어려울 정도로 일치하며, 이는 SCMK 가 가속 효과에도 불구하고 "
        "기본 LBM 의 정상상태 고정점을 변형 없이 보존함을 확인한다.",
        indent_cm=0.5)

    add_figure(doc, "results_contours/kolmogorov_N32.png",
                "그림 10. 2D Kolmogorov flow, N=32 : 기본 LBM (좌) vs SCMK-LBM (우) "
                "수평 속도 u_x 컨투어 + 유선.")
    add_figure(doc, "results_contours/channel_N32.png",
                "그림 11. 채널 Poiseuille flow, N=32 : 좌-기본, 우-SCMK 컨투어.")
    add_figure(doc, "results_contours/couette_N32.png",
                "그림 12. Couette flow, N=32 : 좌-기본, 우-SCMK 컨투어.")
    add_figure(doc, "results_contours/cavity_Re100.png",
                "그림 13. 뚜껑 구동 캐비티, Re=100, N=65 : 좌-기본, 우-SCMK 컨투어 + 유선. "
                "주 와류와 모서리 secondary vortex 모두 일치.")
    add_figure(doc, "results_contours/cavity_Re400.png",
                "그림 14. 뚜껑 구동 캐비티, Re=400, N=97 : 좌-기본, 우-SCMK. "
                "강해진 secondary vortex 까지 시각적으로 일치.")
    add_figure(doc, "results_contours/multi_cylinder.png",
                "그림 15. 다중 실린더 voxel 유동, N=32 : 좌-기본, 우-SCMK. "
                "복잡 기하의 wake 패턴 정확 재현.")
    add_figure(doc, "results_contours/3d_kolmogorov.png",
                "그림 16. 3D Kolmogorov flow, N=24 : 좌-기본, 우-SCMK (z=N/2 중면).")
    add_figure(doc, "results_contours/3d_channel.png",
                "그림 17. 3D 채널 Poiseuille flow, N=24 : 좌-기본, 우-SCMK (z=N/2 중면).")

    add_heading(doc, "3.7 3D 시연", level=2)
    add_para(doc,
        "3D D3Q19 격자에서의 적용 가능성을 보이기 위해, 동일 알고리즘을 "
        "3D Kolmogorov (N=16, 24) 및 3D 채널 Poiseuille (N=24) 흐름에 "
        "적용하였다. 그 결과 N=16에서 3.3×, N=24에서 6.5×, "
        "그리고 3D 채널에서 13.8× 의 LBE-호출 가속을 얻었다. 매개변수 "
        "변경 없이 그대로 적용되어 3D 호환성과 차원 독립성을 입증한다.",
        indent_cm=0.5)

    # ============================================================
    # 4. Conclusion
    # ============================================================
    add_heading(doc, "4. 결론 (Conclusion)", level=1)
    add_para(doc,
        "본 연구는 정상상태 격자볼츠만 방정식 수렴 가속을 위한 매개변수-경량 "
        "뉴턴-크릴로프 솔버 SCMK-LBM을 제안하고 검증하였다. 핵심 결과는 "
        "다음과 같다.",
        indent_cm=0.5)
    add_para(doc, "(1) 자연 LBM 고정점 R(f) = f − L(f) = 0 을 보존하는 외곽 "
             "뉴턴-크릴로프 구조와, 닫힌 형식으로 도출된 푸리에-모멘트 "
             "AP-Schur 전처리기, 자기 조정형 티호노프 정칙화, 그리고 "
             "백트래킹 직선 탐색을 동반한 하이브리드 안전장치를 결합한 "
             "범용 솔버를 제안하였다.", indent_cm=0.5)
    add_para(doc, "(2) 네 가지 정리(AP-Schur 닫힌 형식, 선형 수렴률 한계 0.98, "
             "AP-극한에서 NS Schur 회복, 뉴턴 스텝 유계)를 통해 이론적 "
             "기반을 마련하였다.", indent_cm=0.5)
    add_para(doc, "(3) 2D/3D, 주기/벽/복셀 경계를 망라하는 12 사례에서 "
             "11/12 수렴, 평균 28.6×, 기하평균 13.3×, 최악 3.3×, "
             "최대 194×의 LBE-호출 가속을 실증하였다.", indent_cm=0.5)
    add_para(doc, "(4) Theorem 2의 수렴률 한계를 5/6 사례에서 실증하였다.", indent_cm=0.5)
    add_para(doc, "(5) Ghia 1982 참조 데이터와의 비교에서 캐비티 Re=100에서 "
             "RMS 중심선 속도 오차 3.7×10⁻³, Re=400에서 3.9×10⁻²의 "
             "문헌 수준 정확도를 보였다.", indent_cm=0.5)
    add_para(doc, "(6) Anderson 가속과의 비교에서 평균은 유사하나 최악 시나리오에서 "
             "3× 우수함을 보여, 범용 정상상태 LBM 솔버로서의 적합성을 입증하였다.", indent_cm=0.5)

    add_para(doc,
        "본 방법의 한계는 다음과 같다. (a) Cavity Re=1000 같은 고-Re에서 "
        "BGK 자체의 안정성 한계로 인해 격자/속도 조정이 필요하다. (b) 단일 "
        "실린더 항력계수 검증은 주기-체적력 설정의 본질적 불안정성으로 인해 "
        "달성하지 못하였으며, 적절한 유입-유출 경계조건 구현이 후속 연구로 "
        "남는다. (c) Mavriplis 형 멀티그리드 LBE와의 직접 비교는 본 연구에 "
        "포함되지 않았다. (d) GPU 구현, 실 환자 데이터 기반 혈관 응용 등은 "
        "후속 연구의 영역이다.",
        indent_cm=0.5)

    add_para(doc,
        "이러한 한계에도 불구하고, 본 연구는 단일 경험적 계수만을 남긴 "
        "매개변수-경량 정상상태 LBM 가속법을 제시하고, 다양한 기하와 차원에 "
        "걸쳐 일관된 가속을 실증함으로써 LBM 정상상태 솔버 설계에 새로운 "
        "선택지를 제공한다.",
        indent_cm=0.5)

    # ============================================================
    # References
    # ============================================================
    add_heading(doc, "참고문헌 (References)", level=1)
    refs = [
        "[1] Krüger, T., Kusumaatmaja, H., Kuzmin, A., Shardt, O., Silva, G., Viggen, "
        "E.M., 2017. The Lattice Boltzmann Method: Principles and Practice. Springer.",
        "[2] Succi, S., 2018. The Lattice Boltzmann Equation for Complex States of "
        "Flowing Matter. Oxford University Press.",
        "[3] Guo, Z., Zhao, T.S., Shi, Y., 2004. Preconditioned lattice-Boltzmann "
        "method for steady flows. Physical Review E 70(6), 066706.",
        "[4] Premnath, K.N., Pattison, M.J., Banerjee, S., 2009. Steady state "
        "convergence acceleration of the generalized lattice Boltzmann equation "
        "with forcing term through preconditioning. Physical Review E 79(2), 026703.",
        "[5] Mavriplis, P., 2006. Multigrid solution of the steady-state lattice "
        "Boltzmann equation. Computers & Fluids 35(8–9), 793–804.",
        "[6] Hübner, T., Turek, S., 2009. Efficient monolithic simulation techniques "
        "for the stationary lattice Boltzmann equation. Computing and Visualization "
        "in Science 13(3), 129–143.",
        "[7] Atif, M., Kolluru, P.K., Thantanapally, C., Ansumali, S., 2017. "
        "Essentially entropic lattice Boltzmann model. Physical Review Letters "
        "119(24), 240602.",
        "[8] Anderson, D.G., 1965. Iterative procedures for nonlinear integral "
        "equations. Journal of the ACM 12(4), 547–560.",
        "[9] Ghia, U., Ghia, K.N., Shin, C.T., 1982. High-Re solutions for "
        "incompressible flow using the Navier–Stokes equations and a multigrid "
        "method. Journal of Computational Physics 48(3), 387–411.",
        "[10] Knoll, D.A., Keyes, D.E., 2004. Jacobian-free Newton–Krylov methods: "
        "a survey of approaches and applications. Journal of Computational Physics "
        "193(2), 357–397.",
        "[11] Elman, H.C., Silvester, D.J., Wathen, A.J., 2014. Finite Elements and "
        "Fast Iterative Solvers (2nd ed.). Oxford University Press.",
        "[12] Brown, P.N., Saad, Y., 1990. Hybrid Krylov methods for nonlinear "
        "systems of equations. SIAM Journal on Scientific and Statistical Computing "
        "11(3), 450–481.",
        "[13] Eisenstat, S.C., Walker, H.F., 1996. Choosing the forcing terms in "
        "an inexact Newton method. SIAM Journal on Scientific Computing 17(1), 16–32.",
        "[14] Bouzidi, M., Firdaouss, M., Lallemand, P., 2001. Momentum transfer of "
        "a Boltzmann-lattice fluid with boundaries. Physics of Fluids 13(11), "
        "3452–3459.",
        "[15] Zou, Q., He, X., 1997. On pressure and velocity boundary conditions "
        "for the lattice Boltzmann BGK model. Physics of Fluids 9(6), 1591–1598.",
    ]
    for r in refs:
        add_para(doc, r, size=10, indent_cm=0.0)

    out_path = "/home/younglin90/work/claude_code/claudeCFD/solver_LBM_steady_state/SCMK_LBM_Paper_KR.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    make_paper()

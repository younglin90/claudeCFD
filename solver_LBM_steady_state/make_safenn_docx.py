"""Convert markdown paper to Word docx with Korean styling + figures.

Reads markdown file and renders to docx with Korean fonts, heading styles,
code blocks, tables, and image insertion via ![alt](path) syntax.
"""
import os
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_kor_font(run, name="Malgun Gothic", size=11, bold=False, italic=False, code=False):
    if code:
        name = "Consolas"
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_kor_font(run, size=sizes.get(level, 11), bold=True)
    return p


def add_paragraph(doc, text, *, bold=False, italic=False, size=11, align=None,
                   indent=None, code=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    # Process inline **bold** and `code`
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\$[^$]+\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_kor_font(run, size=size, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_kor_font(run, size=size, code=True)
        elif part.startswith('$') and part.endswith('$'):
            run = p.add_run(part[1:-1])
            set_kor_font(run, size=size, italic=True)
        else:
            run = p.add_run(part)
            set_kor_font(run, size=size, bold=bold, italic=italic, code=code)
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    for line in code_text.split('\n'):
        run = p.add_run(line + '\n')
        set_kor_font(run, size=9, code=True)


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            set_kor_font(run, size=10, bold=(i == 0))


def parse_table_block(lines, idx):
    """Parse markdown table starting at lines[idx]. Returns (rows, end_idx)."""
    rows = []
    while idx < len(lines) and '|' in lines[idx]:
        line = lines[idx].strip()
        if re.match(r'^\|?[\s\-:|]+\|?$', line):  # separator row
            idx += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        idx += 1
    return rows, idx


def md_to_docx(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        text = f.read()
    lines = text.split('\n')
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code fences
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, '\n'.join(code_lines))
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2), level=level)
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|?[\s\-:|]+\|?$', lines[i + 1]):
            rows, i = parse_table_block(lines, i)
            add_table(doc, rows)
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', line):
            i += 1
            continue

        # Blank
        if not line.strip():
            i += 1
            continue

        # Image ![alt](path)
        img_m = re.match(r'^!\[(.*?)\]\(([^)]+)\)\s*$', line.strip())
        if img_m:
            alt, path = img_m.group(1), img_m.group(2)
            if os.path.exists(path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(path, width=Cm(15))
            i += 1
            continue

        # Block math $$...$$
        if line.strip().startswith('$$') and line.strip().endswith('$$') and len(line.strip()) > 4:
            add_paragraph(doc, line.strip()[2:-2], italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        if line.strip() == '$$':
            buf = []
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                buf.append(lines[i])
                i += 1
            i += 1
            add_paragraph(doc, ' '.join(buf).strip(), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            continue

        # Boxed $$\boxed{...}$$
        if '\\boxed{' in line:
            add_paragraph(doc, line.strip(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # Bullet list
        bm = re.match(r'^(\s*)[\-\*]\s+(.*)$', line)
        if bm:
            p = doc.add_paragraph(style='List Bullet')
            indent_level = len(bm.group(1)) // 2
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', bm.group(2))
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2]); set_kor_font(r, bold=True)
                elif part.startswith('`') and part.endswith('`'):
                    r = p.add_run(part[1:-1]); set_kor_font(r, code=True)
                else:
                    r = p.add_run(part); set_kor_font(r)
            i += 1
            continue

        # Numbered list
        nm = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if nm:
            p = doc.add_paragraph(style='List Number')
            r = p.add_run(nm.group(3)); set_kor_font(r)
            i += 1
            continue

        # Regular paragraph
        add_paragraph(doc, line)
        i += 1

    doc.save(docx_path)
    print(f'WROTE {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) >= 3:
        md_to_docx(sys.argv[1], sys.argv[2])
    else:
        md_to_docx('PAPER_SAFENN_KR.md', 'SafeNN_LBM_Paper_KR.docx')

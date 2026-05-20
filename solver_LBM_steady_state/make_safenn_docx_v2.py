"""V2: render LaTeX math to PNG via matplotlib, insert in docx.

Inline $...$  : small PNG inserted as inline shape (height ~0.4cm)
Block $$...$$ : larger PNG centered on its own paragraph (height ~1.0cm)
Boxed         : larger box centered

Renders Korean text + LaTeX math properly. Tables, code blocks, images preserved.
"""
import os
import re
import sys
import hashlib
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MATH_CACHE_DIR = 'math_cache'
os.makedirs(MATH_CACHE_DIR, exist_ok=True)


def render_math(latex, *, block=False, dpi=200):
    """Render LaTeX expr to PNG. Cache by content hash."""
    h = hashlib.md5(f"{latex}|{block}".encode()).hexdigest()[:12]
    path = os.path.join(MATH_CACHE_DIR, f"{h}.png")
    if os.path.exists(path):
        return path
    # Wrap with $...$ for matplotlib mathtext
    if block:
        # Display math: bigger
        fig = plt.figure(figsize=(8, 1.5))
        fig.text(0.5, 0.5, f"${latex}$", fontsize=18,
                  ha='center', va='center')
    else:
        fig = plt.figure(figsize=(4, 0.5))
        fig.text(0.5, 0.5, f"${latex}$", fontsize=12,
                  ha='center', va='center')
    plt.axis('off')
    try:
        plt.savefig(path, dpi=dpi, bbox_inches='tight', pad_inches=0.05,
                     transparent=True)
    except Exception as e:
        # Fallback: render as raw text
        plt.clf()
        fig = plt.figure(figsize=(4, 0.5) if not block else (8, 1.0))
        fig.text(0.5, 0.5, latex, fontsize=12 if not block else 14,
                  ha='center', va='center', family='monospace')
        plt.axis('off')
        plt.savefig(path, dpi=dpi, bbox_inches='tight', pad_inches=0.05,
                     transparent=True)
    plt.close('all')
    return path


def set_kor_font(run, name="Malgun Gothic", size=11, bold=False, italic=False, code=False):
    if code:
        name = "Consolas"
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_kor_font(run, size=sizes.get(level, 11), bold=True)


def add_inline_run(p, text, *, bold=False, italic=False, size=11, code=False):
    """Process text with inline $math$, **bold**, `code`. Render math as PNG."""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\$[^$]+\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_kor_font(run, size=size, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_kor_font(run, size=size, code=True)
        elif part.startswith('$') and part.endswith('$'):
            latex = part[1:-1]
            try:
                png = render_math(latex, block=False)
                run = p.add_run()
                run.add_picture(png, height=Cm(0.45))
            except Exception:
                run = p.add_run(latex)
                set_kor_font(run, size=size, italic=True, code=True)
        else:
            run = p.add_run(part)
            set_kor_font(run, size=size, bold=bold, italic=italic, code=code)


def add_paragraph(doc, text, *, bold=False, italic=False, size=11, align=None,
                   indent=None, code=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    add_inline_run(p, text, bold=bold, italic=italic, size=size, code=code)


def add_block_math(doc, latex):
    """Render block math centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    try:
        png = render_math(latex, block=True)
        run = p.add_run()
        run.add_picture(png, width=Cm(13))
    except Exception:
        run = p.add_run(latex)
        set_kor_font(run, size=11, italic=True, code=True)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    for line in code_text.split('\n'):
        run = p.add_run(line + '\n')
        set_kor_font(run, size=9, code=True)


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            add_inline_run(p, cell_text.strip(), size=10,
                            bold=(i == 0))


def parse_table_block(lines, idx):
    rows = []
    while idx < len(lines) and '|' in lines[idx]:
        line = lines[idx].strip()
        if re.match(r'^\|?[\s\-:|]+\|?$', line):
            idx += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        idx += 1
    return rows, idx


def md_to_docx(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        text = f.read()
    lines = text.split('\n')
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code fences
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, '\n'.join(code_lines))
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2), level=level)
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|?[\s\-:|]+\|?$', lines[i + 1]):
            rows, i = parse_table_block(lines, i)
            add_table(doc, rows)
            continue

        # Image ![alt](path)
        img_m = re.match(r'^!\[(.*?)\]\(([^)]+)\)\s*$', line.strip())
        if img_m:
            path = img_m.group(2)
            if os.path.exists(path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(path, width=Cm(15))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', line):
            i += 1
            continue

        # Blank
        if not line.strip():
            i += 1
            continue

        # Block math $$...$$ on same line
        bm = re.match(r'^\$\$(.+)\$\$\s*$', line.strip())
        if bm:
            latex = bm.group(1)
            # Strip \boxed wrapper if present
            box = re.match(r'\\boxed\{(.+)\}\s*$', latex)
            if box:
                latex = box.group(1)
            add_block_math(doc, latex)
            i += 1
            continue

        # Multiline block math
        if line.strip() == '$$':
            buf = []
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                buf.append(lines[i])
                i += 1
            i += 1
            latex = ' '.join(buf).strip()
            box = re.match(r'\\boxed\{(.+)\}\s*$', latex)
            if box:
                latex = box.group(1)
            add_block_math(doc, latex)
            continue

        # Bullet list
        bl = re.match(r'^(\s*)[\-\*]\s+(.*)$', line)
        if bl:
            p = doc.add_paragraph(style='List Bullet')
            add_inline_run(p, bl.group(2))
            i += 1
            continue

        # Numbered list
        nm = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if nm:
            p = doc.add_paragraph(style='List Number')
            add_inline_run(p, nm.group(3))
            i += 1
            continue

        # Regular paragraph
        add_paragraph(doc, line)
        i += 1

    doc.save(docx_path)
    print(f'WROTE {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) >= 3:
        md_to_docx(sys.argv[1], sys.argv[2])
    else:
        md_to_docx('PAPER_SAFENN_FULL_KR.md', 'SafeNN_LBM_Paper_Full_KR.docx')

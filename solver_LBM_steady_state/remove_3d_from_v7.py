from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화/"
    "SafeNN_LBM_Paper_V7_KR_final_results_augmented.docx"
)
OUT = SRC.with_name("SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d.docx")


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def delete_table(table) -> None:
    element = table._element
    element.getparent().remove(element)


def delete_row(row) -> None:
    element = row._tr
    element.getparent().remove(element)


def main() -> None:
    doc = Document(SRC)

    replacements = {
        "본 장은 baseline Picard-LBM": (
            "본 장은 baseline Picard-LBM, Anderson-accelerated LBM, 그리고 제안 알고리즘 "
            "Safe-NN-SCMK의 비교 결과를 제시한다. §3.1에서는 핵심 6-case의 수렴성"
            "(LBE-call 수, wall-clock, 수렴 이력)을, §3.2에서는 정확도(RMS, L_inf, "
            "relative L2, Ghia centerline deviation, mass drift)를 정량 검증한다. 이어서 "
            "§3.3에서는 추가 검증 case 전체 목록을 기준으로 N=64/128/256 scaling, "
            "cavity Re=400/1000 stress test, backward-facing step, cylinder-wake analogue, "
            "그리고 T-junction mask를 성공/plateau/한계 case로 구분하여 보고한다."
        ),
        "Table 1. 본문 및 확장 검증 전체 case 설정.": (
            "Table 1. 본문 및 확장 검증 전체 case 설정. Case 1-6은 핵심 정량 비교 "
            "benchmark이고, Case 7-15는 major-revision 보완을 위해 추가한 2D scaling, "
            "고-Re stress, mask-flow 검증이다."
        ),
        "위 6 case 외에 Table 1의 Case 7-16": (
            "위 6 case 외에 Table 1의 Case 7-15를 추가 검증했다. 목적은 논문의 평균 "
            "가속률을 높이는 것이 아니라, reviewer가 요구할 가능성이 높은 격자 scaling, "
            "고-Re cavity, 복잡 mask geometry의 경계를 분리해서 보여주는 것이다. 따라서 "
            "아래 결과는 성공 case와 plateau/limitation case를 명시적으로 구분한다."
        ),
        "본 연구는 격자 볼츠만 방법의 정상상태 해를 가속하는 Safe-NN-SCMK를 제안하고": (
            "본 연구는 격자 볼츠만 방법의 정상상태 해를 가속하는 Safe-NN-SCMK를 제안하고, "
            "6 검증 사례와 추가 2D scaling/mask-flow test를 통해 정확도와 수렴성을 점검했다. "
            "본 절은 (1) 핵심 기여의 재정리, (2) prior art와의 정량/구조적 비교, "
            "(3) 결과의 함의와 한계, (4) 향후 작업의 4 부분으로 구성된다."
        ),
        "D2Q9 중심 검증.": (
            "D2Q9 중심 검증. 본문의 정량 결과와 확장 검증은 모두 2D D2Q9 격자에 기반한다. "
            "따라서 본 논문은 2D 정상상태 LBM residual 가속의 가능성과 한계를 다루며, "
            "차원 확장 검증은 본 원고의 claim 범위에 포함하지 않는다."
        ),
    }

    delete_prefixes = [
        "3D production 검증 미완료.",
        "3D D3Q19 / D3Q27 검증.",
    ]

    for p in list(doc.paragraphs):
        text = p.text.strip()
        if any(text.startswith(prefix) for prefix in delete_prefixes):
            delete_paragraph(p)
            continue
        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                set_paragraph_text(p, replacement)
                break

    # Remove D3Q19 row from complete case list.
    for table in doc.tables:
        for row in list(table.rows):
            text = " | ".join(cell.text for cell in row.cells)
            if "D3Q19" in text or "3D " in text or text.startswith("16 |"):
                delete_row(row)

    # Remove an empty table if the 3D smoke-test table lost all data rows.
    for table in list(doc.tables):
        rows_text = [" | ".join(cell.text for cell in row.cells).strip() for row in table.rows]
        if rows_text and rows_text[0].startswith("Case | Picard LBE | Safe LBE"):
            if len(table.rows) == 1:
                delete_table(table)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

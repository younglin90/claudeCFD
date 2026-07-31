"""Remove T-junction content from the active V7 manuscript.

Source: SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_bodyfigures_replaced.docx
Output: SafeNN_LBM_Paper_V7_KR_no_tjunction.docx (sibling of source)

Edits:
  * Drop "T-junction" mentions from body sentences (rewrite phrase, keep flow).
  * Delete Table 1 row 15 (T-junction spec row) and any row in any table that
    contains the literal "T-junction" anywhere in its cells.
  * Delete the dedicated Figure 17 (T-junction field diagnostic): the caption
    paragraph and the immediately preceding paragraph that holds the drawing.
  * Rewrite the Figure 14 caption to drop the T-junction phrase (Backward-step
    + cylinder-wake only). The underlying composite image is left as-is — the
    caption no longer claims a T-junction subpanel.

Run from the repo root.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SRC = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화/"
    "SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_bodyfigures_replaced.docx"
)
DST = SRC.with_name(
    "SafeNN_LBM_Paper_V7_KR_no_tjunction.docx"
)


SENTENCE_REWRITES = [
    (
        "backward-facing step, cylinder-wake analogue, 그리고 T-junction mask를 성공/plateau/한계 case로 구분하여 보고한다.",
        "backward-facing step과 cylinder-wake analogue를 성공/plateau/한계 case로 구분하여 보고한다.",
    ),
    (
        "Figure 14. Additional mask-flow convergence histories. Backward-facing step, cylinder-wake analogue, and T-junction residual histories are included for every additional 2D benchmark.",
        "Figure 14. Additional mask-flow convergence histories for the backward-facing step and cylinder-wake analogue benchmarks.",
    ),
    (
        "정량 해석. Backward-facing step은 16.2x LBE-call 감소, cylinder-wake analogue는 2.43x LBE-call 감소, T-junction은 residual 기준 32.3x 감소를 보였다. Cylinder-wake와 T-junction은 field discrepancy가 더 크므로 정확도 성공 claim이 아니라 복잡 mask에서의 residual/stability stress test로 해석한다.",
        "정량 해석. Backward-facing step은 16.2x LBE-call 감소, cylinder-wake analogue는 2.43x LBE-call 감소를 보였다. Cylinder-wake는 field discrepancy가 더 크므로 정확도 성공 claim이 아니라 복잡 mask에서의 residual/stability stress test로 해석한다.",
    ),
    (
        "Re=1000 cavity의 약한 정확도/느린 wall-clock, T-junction의 큰 relative L2는 현재 알고리즘과 Python 구현의 한계로 본문에 명시한다.",
        "Re=1000 cavity의 약한 정확도/느린 wall-clock은 현재 알고리즘과 Python 구현의 한계로 본문에 명시한다.",
    ),
    (
        "channel 1e-9 plateau, Re=1000 cavity의 약한 정확도와 wall-clock 저하, T-junction의 상대오차 증폭을 드러냈다.",
        "channel 1e-9 plateau와 Re=1000 cavity의 약한 정확도/wall-clock 저하를 드러냈다.",
    ),
]


def paragraph_text(p) -> str:
    return "".join(t.text or "" for t in p._p.iter(qn("w:t")))


def set_paragraph_text(p, new_text: str) -> None:
    runs = p._p.findall(qn("w:r"))
    if not runs:
        return
    # Clear all but first run; put all text into first run's first <w:t>.
    first_run = runs[0]
    first_t = first_run.find(qn("w:t"))
    if first_t is None:
        first_t = first_run.makeelement(qn("w:t"), {})
        first_run.append(first_t)
    first_t.text = new_text
    # Remove additional <w:t> in first run
    for extra in first_run.findall(qn("w:t"))[1:]:
        first_run.remove(extra)
    # Remove subsequent runs entirely
    for extra_run in runs[1:]:
        p._p.remove(extra_run)


def rewrite_body(doc) -> int:
    n = 0
    for p in doc.paragraphs:
        txt = paragraph_text(p)
        for old, new in SENTENCE_REWRITES:
            if old in txt:
                set_paragraph_text(p, txt.replace(old, new))
                n += 1
                txt = paragraph_text(p)
    return n


def find_paragraph_with(doc, needle: str):
    for p in doc.paragraphs:
        if needle in paragraph_text(p):
            return p
    return None


def remove_figure17(doc) -> bool:
    caption = find_paragraph_with(doc, "Figure 17. T-junction field diagnostic")
    if caption is None:
        return False
    parent = caption._p.getparent()
    siblings = list(parent)
    idx = siblings.index(caption._p)
    # The preceding paragraph (or two) holds the drawing. Remove any preceding
    # paragraphs that contain a <w:drawing> until we hit a non-drawing paragraph.
    removed = [caption._p]
    parent.remove(caption._p)
    j = idx - 1
    while j >= 0:
        sib = siblings[j]
        if sib.tag != qn("w:p"):
            break
        if sib.find(".//" + qn("w:drawing")) is None:
            break
        removed.append(sib)
        parent.remove(sib)
        j -= 1
    return True


def remove_table_rows_with_tjunction(doc) -> int:
    n_removed = 0
    for tbl in doc.tables:
        rows_to_remove = []
        for row in tbl.rows:
            row_text = " ".join(c.text for c in row.cells)
            if "T-junction" in row_text:
                rows_to_remove.append(row)
        for row in rows_to_remove:
            row._tr.getparent().remove(row._tr)
            n_removed += 1
    return n_removed


def count_tjunction(doc) -> int:
    n = 0
    for p in doc.paragraphs:
        n += paragraph_text(p).count("T-junction")
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                n += cell.text.count("T-junction")
    return n


def main():
    if not SRC.exists():
        raise SystemExit(f"missing source docx: {SRC}")
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))
    print(f"initial T-junction mentions: {count_tjunction(doc)}")
    n_body = rewrite_body(doc)
    print(f"body sentences rewritten: {n_body}")
    fig17 = remove_figure17(doc)
    print(f"figure 17 removed: {fig17}")
    n_rows = remove_table_rows_with_tjunction(doc)
    print(f"table rows removed: {n_rows}")
    doc.save(str(DST))
    # Re-open to confirm and recount
    doc2 = Document(str(DST))
    remaining = count_tjunction(doc2)
    print(f"remaining T-junction mentions after edit: {remaining}")
    print(f"wrote: {DST}")


if __name__ == "__main__":
    main()

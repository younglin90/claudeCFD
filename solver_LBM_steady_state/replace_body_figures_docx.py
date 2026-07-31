from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ONEDRIVE_DIR = Path(
    "/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화"
)
SOURCE = ONEDRIVE_DIR / "SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_cavityfixed.docx"
OUTPUT = ONEDRIVE_DIR / "SafeNN_LBM_Paper_V7_KR_final_results_augmented_no3d_bodyfigures_replaced.docx"
FIG = Path("paper_revision_data/figures_complete")


REPLACEMENTS = [
    (
        177,
        FIG / "figU1_core6_convergence_updated.png",
        6.8,
        "Figure 1. Complete native-residual convergence histories for the six core 2D benchmarks. "
        "Picard-LBM is shown by dashed gray curves and Safe-NN-SCMK by solid red curves. "
        "The Re=400 cavity row is retained as a stiff coarse-grid stress test rather than an oscillation-free field-validation figure.",
    ),
    (
        180,
        FIG / "figU11_all_2d_cases_speedup_status_updated.png",
        6.9,
        "Figure 2. LBE-call and wall-clock speedup/status summary for all 2D validation cases. "
        "Successful convergence, strict-tolerance plateau behavior, and high-Re limitations are separated explicitly.",
    ),
    (
        195,
        FIG / "figU2_core_profiles_updated.png",
        6.6,
        "Figure 3. Analytical-profile validation for the smooth core benchmarks. "
        "Kolmogorov, channel, and Couette profiles compare Picard and Safe-NN against their analytical steady profiles.",
    ),
    (
        198,
        FIG / "figU3_cavity_centerlines_updated.png",
        6.5,
        "Figure 4. Lid-driven cavity centerline validation. "
        "Re=100 and Re=400 centerlines are compared with Ghia et al. reference data; "
        "the Re=400 N=49 result is interpreted as a coarse-grid stress diagnostic.",
    ),
    (
        201,
        FIG / "figU4_multicylinder_field_updated.png",
        6.7,
        "Figure 5. Multi-cylinder voxel-mask field validation. "
        "Picard velocity magnitude, Safe-NN velocity magnitude, and absolute difference are shown for the core obstacle benchmark.",
    ),
    (
        204,
        FIG / "figU8_cavity_re400_n65_polished_field_updated.png",
        6.7,
        "Figure 6. Oscillation-free lid-driven cavity Re=400, N=65 field after final Picard polish. "
        "The Safe-NN field is post-relaxed to the tight residual level and compared with the tight Picard baseline.",
    ),
    (
        207,
        FIG / "figU3b_cavity_re400_n49_polished_centerline_updated.png",
        6.4,
        "Figure 7. Cavity Re=400, N=49 post-smoothed centerline diagnostic. "
        "A long post-smoothing run reduces the residual only to about 2.5e-7, so this coarse-grid case is retained as a limitation/stress diagnostic.",
    ),
    (
        210,
        FIG / "figU_kolmogorov_N64_field_updated.png",
        6.7,
        "Figure 8. Kolmogorov N=64 field validation. "
        "Picard/Safe-NN fields and pointwise difference show phase-consistent agreement across the full domain.",
    ),
    (
        213,
        FIG / "figU_channel_N64_field_updated.png",
        6.7,
        "Figure 9. Channel N=64 field validation. "
        "Picard/Safe-NN fields and pointwise difference show that the wall-bounded profile is not spatially biased by the accelerator.",
    ),
    (
        216,
        FIG / "figS2_mass_conservation_updated.png",
        6.3,
        "Figure 10. Mass-conservation check for the core benchmarks. "
        "Relative mass drift remains small for Picard and Safe-NN, with the cavity cases reported separately because their wall forcing is stiffer.",
    ),
    (
        222,
        FIG / "figU5_scaling_convergence_updated.png",
        6.8,
        "Figure 11. Grid-scaling convergence histories. "
        "Kolmogorov and channel cases at N=64, 128, and 256 are shown; channel N=128/256 expose strict-tolerance plateau behavior.",
    ),
    (
        224,
        FIG / "figU6_scaling_speedups_updated.png",
        6.6,
        "Figure 12. Grid-scaling speedup/status summary. "
        "LBE-call and wall-clock speedups are reported together, with plateau cases retained as limitations rather than success claims.",
    ),
    (
        226,
        FIG / "figU7_high_re_cavity_convergence_updated.png",
        6.5,
        "Figure 13. High-Re cavity convergence histories. "
        "Re=400 N=65 uses final Picard polish; Re=1000 N=129 is reported as a high-Re limitation case with line-search safeguard.",
    ),
    (
        228,
        FIG / "figU9_extra_mask_convergence_updated.png",
        6.8,
        "Figure 14. Additional mask-flow convergence histories. "
        "Backward-facing step, cylinder-wake analogue, and T-junction residual histories are included for every additional 2D benchmark.",
    ),
    (
        230,
        FIG / "figU10_backward_step_field_updated.png",
        6.7,
        "Figure 15. Backward-facing step field validation. "
        "Picard velocity magnitude, Safe-NN velocity magnitude, and absolute difference are shown.",
    ),
    (
        234,
        FIG / "figU10_cylinder_wake_field_updated.png",
        6.7,
        "Figure 16. Cylinder-wake analogue field validation. "
        "This is a steady obstacle benchmark and is not used as a true unsteady vortex-shedding validation.",
    ),
    (
        236,
        FIG / "figU10_t_junction_field_updated.png",
        6.7,
        "Figure 17. T-junction field diagnostic. "
        "The T-junction is reported as a narrow-mask residual stress test; its field discrepancy is treated as a limitation.",
    ),
    (
        239,
        FIG / "figU12_safeguard_diagnostics_updated.png",
        6.9,
        "Figure 18. Safeguard and fallback activity. "
        "Lookahead rejection, residual restart, line-search rejection, and final-polish activity show where the safety layers are active.",
    ),
    (
        242,
        FIG / "figS1_residual_error_scatter_updated.png",
        5.8,
        "Figure 19. Residual-error relation for limitation cases. "
        "Final residuals are plotted against relative L2 velocity differences where field discrepancy is central to interpretation.",
    ),
    (
        244,
        FIG / "figS3_cavity_final_polish_diagnostic_updated.png",
        5.9,
        "Figure 20. Cavity final-polish diagnostic. "
        "Residual and relative L2 error decrease during Picard final-polish for Re=400 cavity, identifying the origin of the earlier contour oscillation.",
    ),
]


PARAGRAPH_UPDATES = {
    176: (
        "Figure 1 replaces the earlier partial convergence plot with complete residual histories for all six core 2D benchmarks. "
        "The stiff Re=400 cavity is shown with the production tolerance history; the tightened/post-polish behavior is separated into Figures 6, 7, and 20."
    ),
    179: (
        "Figure 2 summarizes LBE-call and wall-clock behavior across the full 2D validation list, not only the original six cases. "
        "The plot separates successful convergence from strict-tolerance plateau and high-Re limitation cases."
    ),
    187: (
        "stiff regime (cavity Re=400): the production/stress-test result at N=49 and tolerance 5e-7 gives a useful LBE-call reduction, "
        "but the oscillation-free verification figure requires N=65 plus final Picard polish to 5e-8. Consequently this case is presented "
        "as a stability and limitation result rather than as an absolute wall-clock acceleration claim."
    ),
    194: (
        "Figure 3 combines the smooth analytical-profile validations. Kolmogorov, channel, and Couette all preserve the expected analytical shape, "
        "showing that the accelerator does not introduce an additional macroscopic-profile bias in the smooth benchmark class."
    ),
    197: (
        "Figure 4 reports lid-driven cavity centerline validation against Ghia et al. Re=100 and Re=400 data. "
        "The Re=400 N=49 curve is retained as a coarse-grid stress diagnostic; the polished field validation is reported separately in Figure 6."
    ),
    200: (
        "Figure 5 shows the multi-cylinder voxel-mask field comparison. This obstacle-dominated case is less favorable than the smooth periodic cases, "
        "but the Safe-NN field remains aligned with the Picard-LBM reference at the reported tolerance."
    ),
    203: (
        "Figure 6 replaces the earlier oscillatory Re=400 cavity contour with the N=65 final-polish result. "
        "The post-relaxed Safe-NN field reaches residual 4.53e-8 and relative L2 difference 2.88e-3 against the tight Picard baseline."
    ),
    206: (
        "Figure 7 explains why the old N=49 Re=400 history should not be used as the final polished contour evidence. "
        "Even after long post-smoothing, the coarse-grid residual remains near 2.5e-7, so the N=49 run is a stress diagnostic."
    ),
    209: (
        "Figures 8-10 add full-field and conservation checks that support the accuracy discussion beyond centerline profiles."
    ),
    212: (
        "Figure 8 shows the N=64 Kolmogorov field comparison, including the pointwise Safe-NN/Picard difference."
    ),
    215: (
        "Figure 9 shows the N=64 channel field comparison, including the pointwise Safe-NN/Picard difference."
    ),
    221: (
        "Figures 11-20 report the major-revision extension set: N-scaling, high-Re cavity, additional 2D mask-flow benchmarks, safeguard diagnostics, "
        "mass/residual-error checks, and the cavity final-polish diagnostic."
    ),
    232: (
        "정량 해석. Backward-facing step은 16.2x LBE-call 감소, cylinder-wake analogue는 2.43x LBE-call 감소, "
        "T-junction은 residual 기준 32.3x 감소를 보였다. Cylinder-wake와 T-junction은 field discrepancy가 더 크므로 "
        "정확도 성공 claim이 아니라 복잡 mask에서의 residual/stability stress test로 해석한다."
    ),
    238: (
        "정량 해석. Kolmogorov는 N=128/256에서도 강한 LBE-call scaling을 유지했다. Channel N=128/256은 "
        "strict 1e-9 tolerance에서 residual plateau를 보였으므로 limitation으로 표시했다. Re=400 cavity는 N=65 final polish 후 "
        "residual 4.53e-8, relative L2 2.88e-3까지 개선되며, Re=1000 N=129는 고-Re 한계 case로 보고한다."
    ),
    241: (
        "Figures 18-20 collect the reviewer-facing diagnostics: safeguard activation, residual-error relation, and final-polish effect."
    ),
}


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def replace_picture(paragraph, path: Path, width: float) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def main() -> None:
    doc = Document(SOURCE)
    for idx, text in PARAGRAPH_UPDATES.items():
        replace_paragraph_text(doc.paragraphs[idx], text)

    for idx, path, width, caption in REPLACEMENTS:
        replace_picture(doc.paragraphs[idx], path, width)
        replace_paragraph_text(doc.paragraphs[idx + 1], caption)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "papers_data" / "ap_schur_only_methods_section.docx"
SUMMARY = ROOT / "papers_data" / "summary_latest_ap_schur_only_proposed.csv"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY = RGBColor(88, 88, 88)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "F2F4F7"
BORDER = "D8DEE8"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=BORDER, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "3")
        node.set(qn("w:color"), color)


def add_para(doc, text="", style=None, size=11, color=None, bold=False, italic=False, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_equation(doc, equation, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    set_paragraph_shading(p, LIGHT_FILL)
    set_paragraph_border(p)
    r = p.add_run(equation)
    set_run_font(r, name="Cambria Math", size=11.5, color=INK)
    if note:
        r2 = p.add_run("\n" + note)
        set_run_font(r2, size=9.5, color=GRAY, italic=True)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], TABLE_FILL)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if i == 0:
                set_cell_shading(cells[i], "FBFCFE")
            r = p.add_run(str(value))
            set_run_font(r, size=9.2, color=RGBColor(30, 30, 30))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def load_case_counts():
    if not SUMMARY.exists():
        return "27 proposed AP-Schur-only rows were expected in papers_data, but the summary CSV was not found."
    with SUMMARY.open(newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    levels = {}
    variants = {}
    for row in rows:
        levels[row.get("scaling_level", "")] = levels.get(row.get("scaling_level", ""), 0) + 1
        variants[row.get("method_variant", "")] = variants.get(row.get("method_variant", ""), 0) + 1
    level_text = ", ".join(f"{k}x: {v}" for k, v in sorted(levels.items()))
    variant_text = ", ".join(f"{k}: {v}" for k, v in sorted(variants.items()))
    return f"Current archive check: {len(rows)} AP-Schur-only proposed rows in papers_data ({level_text}); variants recorded as {variant_text}."


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Calibri"
        styles[name]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        styles[name]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("Methods draft | AP-Schur-only steady LBM solver")
    set_run_font(r, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Prepared for manuscript development")
    set_run_font(r, size=9, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Methods: Geometry-Aware AP-Schur-Only Solver for Steady Lattice Boltzmann Problems")
    set_run_font(r, size=22, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("Draft Methods section for a top-tier SCI manuscript")
    set_run_font(r, size=12.5, color=GRAY, italic=True)

    meta = [
        ("Proposed method", "AP-Schur-only native-residual JFNK acceleration for steady LBM"),
        ("Excluded from final proposed method", "RRE and case-specific tuning coefficients"),
        ("Convergence basis", "Macroscopic L2 residual plus residual-floor/plateau confirmation"),
        ("Archive note", load_case_counts()),
    ]
    add_simple_table(doc, ["Item", "Specification"], meta, [2300, 7060])

    add_heading(doc, "2. Methods", 1)
    add_para(
        doc,
        "This section describes the proposed AP-Schur-only solver as a nonlinear acceleration method for steady-state lattice Boltzmann computations. "
        "The central design principle is deliberately conservative: the discretized LBM operator, the boundary-condition implementation, and the benchmark reference comparisons are not altered by the accelerator. "
        "Instead, AP-Schur generates residual-reducing candidate corrections for the same native steady LBM equations. A candidate is accepted only when it decreases the common macroscopic residual and satisfies admissibility checks. "
        "Therefore, the method should be interpreted as a preconditioned nonlinear iteration for the steady LBM residual, not as a different physical model or as a post-processing correction.",
    )

    add_heading(doc, "2.1 Steady LBM Formulation", 2)
    add_para(
        doc,
        "Let Ω_f denote the set of fluid nodes after excluding solid cells, masks, and immersed obstacles. At each fluid node x and lattice direction i = 0,...,q-1, the LBM stores a distribution f_i(x). "
        "The discrete velocity set is denoted by c_i, the quadrature weights by w_i, and the lattice sound speed by c_s. The macroscopic density, momentum, and pressure are obtained from velocity moments of f:",
    )
    add_equation(doc, "ρ(x) = Σᵢ fᵢ(x),        ρ(x)u(x) = Σᵢ cᵢ fᵢ(x),        p(x) = cₛ² ρ(x).")
    add_para(
        doc,
        "For the single-relaxation-time formulation used in the benchmark solver, the collision step relaxes f toward the second-order equilibrium distribution:",
    )
    add_equation(
        doc,
        "fᵢ*(x) = fᵢ(x) - ω [ fᵢ(x) - fᵢᵉᵠ(ρ,u) ],",
    )
    add_equation(
        doc,
        "fᵢᵉᵠ = wᵢρ [ 1 + (cᵢ·u)/cₛ² + (cᵢ·u)²/(2cₛ⁴) - (u·u)/(2cₛ²) ].",
        "For D2Q9, c_s² = 1/3 and ν = c_s²(1/ω - 1/2) in lattice units.",
    )
    add_para(
        doc,
        "Streaming and boundary treatment are then applied by the native LBM update operator. We write one complete collide-stream-boundary update compactly as G(f). "
        "The steady state is the fixed point of this native operator:",
    )
    add_equation(doc, "R(f) = G(f) - f = 0.")
    add_para(
        doc,
        "The residual R(f) is called native because it is evaluated by the same LBM update, masks, wall rules, inlet/outlet treatment, and obstacle geometry used by the reference Picard iteration. "
        "This native-residual definition is important for fairness: every solver is judged against the same discrete steady equation.",
    )

    add_heading(doc, "2.2 Macroscopic L2 Residual Used for Convergence", 2)
    add_para(
        doc,
        "Although R(f) is defined in distribution space, convergence is monitored in macroscopic variables because the steady-flow target is a stable pressure-velocity field. "
        "For two consecutive accepted states f^k and f^{k+1}, the residual components are measured over fluid nodes only:",
    )
    add_equation(
        doc,
        "rₚᵏ = ||pᵏ⁺¹ - pᵏ||₂ / √|Ω_f|,     rᵤᵏ = ||uₓᵏ⁺¹ - uₓᵏ||₂ / √|Ω_f|,     rᵥᵏ = ||uᵧᵏ⁺¹ - uᵧᵏ||₂ / √|Ω_f|.",
    )
    add_equation(
        doc,
        "r_macroᵏ = √[ (rₚᵏ)² + (rᵤᵏ)² + (rᵥᵏ)² + (r_wᵏ)² ],",
        "The z-velocity term r_w is included for three-dimensional extensions and is zero for the present two-dimensional cases.",
    )
    add_para(
        doc,
        "The initial-relative residual is also recorded as r_macro^k / r_macro^0. The historical f-RMS residual is retained in the output files as an auxiliary diagnostic, but it is not the primary convergence quantity in the final protocol. "
        "Using a macroscopic L2 norm avoids a misleading situation in which microscopic distribution changes are small while the pressure or velocity field is still drifting.",
    )

    add_heading(doc, "2.3 AP-Schur Idea", 2)
    add_para(
        doc,
        "The proposed accelerator is based on a moment-space approximation to the Schur complement of the native residual Jacobian. Near a steady state, Newton's method would seek a correction δf satisfying:",
    )
    add_equation(doc, "J_f(fᵏ) δf = -R(fᵏ),        J_f = ∂R/∂f.")
    add_para(
        doc,
        "Directly forming J_f is not practical for masked geometries and LBM boundary operators. The key observation is that slow steady-state error is dominated by hydrodynamic, low-frequency moment modes, whereas kinetic modes are strongly damped by collision and native LBM smoothing. "
        "The AP-Schur step therefore projects the residual equation to a compact moment space, solves an approximate hydrodynamic correction, then lifts that correction back to distribution space.",
    )
    add_equation(
        doc,
        "m = P f,        Sₘ ≈ P J_f P†,        Sₘ δm ≈ -P R(fᵏ),        δf_AP = P†δm.",
        "P maps distributions to conserved and near-conserved moments; P† is a consistent lifting from moments back to distributions.",
    )
    add_para(
        doc,
        "Here 'AP' is used in the asymptotic-preserving sense: the preconditioner is designed to remain effective when the macroscopic steady flow is governed by diffusive or incompressible slow modes rather than by the raw collide-stream relaxation rate. "
        "The Schur approximation targets the part of the residual that standard Picard LBM removes slowly. The native LBM update is still responsible for enforcing detailed kinetic relaxation and boundary consistency.",
    )

    add_heading(doc, "2.4 Jacobian-Free Correction", 2)
    add_para(
        doc,
        "The implementation uses Jacobian-free products so that the solver does not need to assemble or store a global Jacobian. For a trial direction v, the native residual Jacobian action is approximated by a finite difference:",
    )
    add_equation(doc, "J_f(f)v ≈ [ R(f + εv) - R(f) ] / ε,        ε = √ε_machine (1 + ||f||₂) / ||v||₂.")
    add_para(
        doc,
        "A short Krylov solve obtains the moment-preconditioned correction. The AP-Schur operation acts as a left preconditioner for the slow macroscopic part of the problem, while the residual evaluation remains the original LBM operator. "
        "This combination gives two practical advantages: the memory footprint stays close to the native solver, and the preconditioner remains compatible with complex masks and boundary rules because all nonlinear residual evaluations pass through G(f).",
    )

    add_heading(doc, "2.5 Candidate Generation and Acceptance", 2)
    add_para(
        doc,
        "At each outer cycle, the solver compares a conservative native candidate with an AP-Schur candidate. The AP-Schur correction is never applied blindly. Instead, the method performs damped candidate trials:",
    )
    add_equation(doc, "f_trial(α) = fᵏ + α δf_AP,        α ∈ {1, 1/2, 1/4, ...}.")
    add_para(
        doc,
        "Each trial is passed through the same LBM boundary and settling operations used by the baseline solver, then evaluated with the common macroscopic residual. "
        "A trial can be accepted only if it improves the residual and remains physically admissible. If no AP-Schur candidate passes the acceptance tests, the method falls back to the native LBM candidate for that cycle. This fallback is essential: AP-Schur is an accelerator, not a replacement for the stable native operator.",
    )
    add_simple_table(
        doc,
        ["Acceptance gate", "Purpose"],
        [
            ("Finite field check", "Reject NaN/Inf pressure, density, velocity, or distribution values."),
            ("Positive density branch", "Reject corrections that push the density field toward a nonphysical low-density branch."),
            ("Residual monotonicity", "Accept AP-Schur only when r_macro decreases relative to the competing native candidate or the last accepted state."),
            ("Boundary consistency", "Reapply native wall, inlet, outlet, and mask operators before residual evaluation."),
            ("No reference-data injection", "Analytic, Ghia, or tight-reference values are never used inside the solve; they are post-run error measures only."),
        ],
        [2300, 7060],
    )

    add_heading(doc, "2.6 Geometry-Aware Treatment", 2)
    add_para(
        doc,
        "Complex CFD benchmarks contain rectangular channels, closed cavities, backward-facing steps, cylinder masks, multi-cylinder arrays, and T-junction open boundaries. "
        "The proposed method uses a single geometry-aware AP-Schur-only algorithm for all of them. Geometry awareness enters through three mechanisms. First, all norms and moment projections are restricted to Ω_f, so solid nodes never contribute to the correction norm or residual. "
        "Second, the correction is evaluated only through the native operator G(f), meaning wall bounce-back, inlet/outlet reconstruction, and mask handling remain centralized in one implementation. "
        "Third, open-boundary branches are checked for physically meaningful density and velocity fields before acceptance. These safeguards prevent the Schur correction from exploiting pressure-gauge freedom or open-boundary degrees of freedom to reduce a residual numerically while degrading the physical solution.",
    )

    add_heading(doc, "2.7 Convergence and Plateau Criterion", 2)
    add_para(
        doc,
        "A run is considered converged only when both a residual threshold and a residual-floor condition are satisfied. The absolute criterion is scaled by the benchmark tolerance:",
    )
    add_equation(doc, "r_macroᵏ < 5τ.")
    add_para(
        doc,
        "The plateau test prevents premature termination immediately after a rapid AP-Schur drop. Let W be the plateau window and define the recent fractional improvement as:",
    )
    add_equation(doc, "I_Wᵏ = [ r_macroᵏ⁻ᵂ - r_macroᵏ ] / max(r_macroᵏ⁻ᵂ, ε_floor).")
    add_para(
        doc,
        "The floor condition is satisfied when the recent improvement is at most 5%, after the required minimum number of LBM evaluations has elapsed. In words, the residual must be small and must also have reached the asymptotic floor of the current discrete problem. "
        "If a saved run satisfies the residual threshold but not the plateau condition, continuation is performed from the saved state using the identical method and the identical convergence test. Continuation changes the iteration count, not the algorithm.",
    )

    add_heading(doc, "2.8 Algorithm", 2)
    add_simple_table(
        doc,
        ["Step", "Operation"],
        [
            ("1", "Initialize f from the benchmark initial condition and apply the native boundary operator."),
            ("2", "Evaluate p, u, v, and the initial macroscopic L2 residual r_macro^0 over fluid cells."),
            ("3", "Generate a conservative native LBM candidate using the same collide-stream-boundary operator G(f) as the baseline Picard solver."),
            ("4", "Project the native residual to moment space, approximately solve the AP-Schur moment correction, and lift the correction to distribution space."),
            ("5", "Test damped AP-Schur candidates. Reapply boundary rules and compute the common macroscopic residual for every trial."),
            ("6", "Accept the AP-Schur candidate only if all admissibility gates pass and the residual improves; otherwise accept the native candidate."),
            ("7", "Record wall time, LBE calls, macro residual components, f-RMS diagnostic residual, AP-Schur trials, and AP-Schur accepts."),
            ("8", "Stop only when r_macro < 5τ and the residual plateau/floor criterion is satisfied. Otherwise continue from the latest accepted state."),
        ],
        [900, 8460],
    )

    add_heading(doc, "2.9 Fair Benchmark Protocol", 2)
    add_para(
        doc,
        "The proposed solver is evaluated under the same residual definition, stopping rule, mesh scaling, and post-processing metrics as the reference methods. "
        "The final proposed method uses AP-Schur only; RRE is disabled. This choice was made after ablation showed that AP-Schur alone gave the best wall-clock behavior while preserving the same accuracy class. "
        "No benchmark-specific tuning coefficient is introduced. Parameters such as damping candidates, residual thresholds, plateau tolerance, and admissibility gates are global method settings rather than case-dependent knobs.",
    )
    add_simple_table(
        doc,
        ["Benchmark family", "Reference comparison used after the solve"],
        [
            ("Channel Poiseuille", "Analytic velocity profile and pressure-driven flow diagnostics."),
            ("Couette flow", "Analytic linear velocity profile."),
            ("Lid-driven cavity Re100/Re400/Re1000", "Ghia et al. centerline velocity profiles and field-level consistency metrics."),
            ("Backward-facing step", "Tight-reference or previously converged numerical field and flow-feature diagnostics."),
            ("Cylinder wake Re40", "Reference steady wake field and wake-related diagnostics."),
            ("Six-cylinder mask", "Reference masked-domain velocity/pressure field."),
            ("T-junction", "Reference open-boundary junction field and boundary-consistency diagnostics."),
        ],
        [2500, 6860],
    )

    add_heading(doc, "2.10 Why AP-Schur Accelerates the Solver", 2)
    add_para(
        doc,
        "Steady LBM Picard iteration is robust because each step is a physical collide-stream update, but it can be slow when the remaining error is a long-wavelength hydrodynamic mode. "
        "Such modes have weak damping under local relaxation and require many lattice sweeps to equilibrate globally. The AP-Schur preconditioner attacks this bottleneck directly: it approximates the macroscopic pressure-velocity Schur response and proposes a global correction in the space where the slow error lives. "
        "Kinetic and boundary-local errors are then controlled by the native update and by residual-monotone acceptance. The resulting method combines a global hydrodynamic correction with local LBM stability.",
    )
    add_equation(
        doc,
        "error = slow hydrodynamic moment component + fast kinetic component;     AP-Schur targets the first, native LBM damps the second.",
    )

    add_heading(doc, "2.11 Reproducibility Notes", 2)
    add_para(
        doc,
        "All saved histories include the accepted macroscopic residual sequence, wall time, LBE-call count, final residual components, and the auxiliary f-RMS residual. "
        "The proposed rows in the current archive are labeled as uniform_ap_schur_only, with continued labels used only for cases restarted from saved states to satisfy the same plateau criterion. "
        "These labels document the computational pathway without changing the mathematical method.",
    )
    add_para(
        doc,
        "The method should be presented in the manuscript as an acceleration and preconditioning strategy for the native steady LBM residual. Accuracy claims should be tied to the same discretization, boundary treatment, Mach-number setting, and mesh refinement used by the reference solvers. "
        "The appropriate novelty claim is therefore not that AP-Schur changes the physical steady equation, but that it provides a geometry-aware, residual-safe, and case-uniform Schur preconditioning route to reach that equation substantially faster.",
    )

    add_heading(doc, "Suggested In-Text Citations", 2)
    add_para(
        doc,
        "The final manuscript should cite the standard LBM equilibrium/discretization literature, the Ghia et al. lid-driven cavity benchmark for Re100/Re400/Re1000 centerlines, GMRES and Jacobian-free Newton-Krylov literature, and Schur-complement/preconditioning references. "
        "Exact bibliography formatting should be completed during manuscript assembly.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())

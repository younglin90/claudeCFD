from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "papers_data" / "ap_schur_only_methods_section_ko.docx"
SUMMARY = ROOT / "papers_data" / "summary_latest_ap_schur_only_proposed.csv"

FONT_KO = "Noto Sans CJK KR"
FONT_MATH = "Cambria Math"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY = RGBColor(88, 88, 88)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "F2F4F7"
BORDER = "D8DEE8"


def set_run_font(run, name=FONT_KO, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    lang = run._element.rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.rPr.append(lang)
    lang.set(qn("w:val"), "ko-KR")
    lang.set(qn("w:eastAsia"), "ko-KR")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=BORDER, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "3")
        node.set(qn("w:color"), color)


def add_para(doc, text="", size=10.5, color=None, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.28
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_equation(doc, equation, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    set_paragraph_shading(p, LIGHT_FILL)
    set_paragraph_border(p)
    r = p.add_run(equation)
    set_run_font(r, name=FONT_MATH, size=11.5, color=INK)
    if note:
        r2 = p.add_run("\n" + note)
        set_run_font(r2, size=9, color=GRAY, italic=True)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], TABLE_FILL)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9.2, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if i == 0:
                set_cell_shading(cells[i], "FBFCFE")
            r = p.add_run(str(value))
            set_run_font(r, size=8.8, color=RGBColor(30, 30, 30))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def load_case_counts():
    if not SUMMARY.exists():
        return "papers_data의 최신 AP-Schur-only summary CSV를 찾지 못했습니다."
    with SUMMARY.open(newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    levels = {}
    variants = {}
    for row in rows:
        levels[row.get("scaling_level", "")] = levels.get(row.get("scaling_level", ""), 0) + 1
        variants[row.get("method_variant", "")] = variants.get(row.get("method_variant", ""), 0) + 1
    level_text = ", ".join(f"{k}x: {v}" for k, v in sorted(levels.items()))
    variant_text = ", ".join(f"{k}: {v}" for k, v in sorted(variants.items()))
    return f"현재 보관 결과: papers_data에 AP-Schur-only proposed row {len(rows)}개가 있음({level_text}); 기록된 variant는 {variant_text}."


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_KO
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_KO)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_KO)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.28

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = FONT_KO
        styles[name]._element.rPr.rFonts.set(qn("w:ascii"), FONT_KO)
        styles[name]._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_KO)
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("Methods draft | AP-Schur-only steady LBM solver")
    set_run_font(r, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("논문 Methods 초안")
    set_run_font(r, size=9, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Methods: 정상상태 LBM을 위한 Geometry-Aware AP-Schur-Only Solver")
    set_run_font(r, size=20, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("Top-tier SCI 논문 작성을 위한 한글 Methods 상세 초안")
    set_run_font(r, size=12, color=GRAY, italic=True)

    add_simple_table(
        doc,
        ["항목", "명세"],
        [
            ("제안 방법", "steady LBM native residual을 대상으로 하는 AP-Schur-only JFNK 가속 solver"),
            ("최종 제안법에서 제외", "RRE 및 케이스별 튜닝 계수"),
            ("수렴 판정", "macroscopic L2 residual과 residual floor/plateau 조건을 동시에 만족"),
            ("결과 보관 상태", load_case_counts()),
        ],
        [2100, 7260],
    )

    add_heading(doc, "2. Methods", 1)
    add_para(
        doc,
        "본 절에서는 제안 방법인 AP-Schur-only solver를 정상상태 lattice Boltzmann method(LBM)를 위한 비선형 수렴 가속법으로 설명한다. "
        "핵심 원칙은 보수적으로 설정하였다. 즉, 이 방법은 LBM의 이산화 방정식, 경계조건 구현, benchmark reference 비교 방식을 바꾸지 않는다. "
        "AP-Schur는 동일한 native steady LBM residual을 더 빠르게 감소시키는 후보 correction을 생성할 뿐이며, residual 감소와 물리적 admissibility 조건을 모두 만족할 때에만 correction을 받아들인다. "
        "따라서 제안법은 새로운 유동 물리모델이 아니라, 동일한 정상상태 LBM 방정식에 도달하기 위한 preconditioned nonlinear iteration으로 해석되어야 한다.",
    )

    add_heading(doc, "2.1 정상상태 LBM 정식화", 2)
    add_para(
        doc,
        "Ω_f를 solid cell, mask, immersed obstacle을 제외한 fluid node 집합이라고 하자. 각 fluid node x와 lattice direction i = 0,...,q-1에 대해 LBM은 distribution function fᵢ(x)를 저장한다. "
        "discrete velocity는 cᵢ, quadrature weight는 wᵢ, lattice sound speed는 cₛ로 표기한다. density, momentum, pressure는 f의 velocity moment로부터 계산된다.",
    )
    add_equation(doc, "ρ(x) = Σᵢ fᵢ(x),        ρ(x)u(x) = Σᵢ cᵢ fᵢ(x),        p(x) = cₛ²ρ(x).")
    add_para(
        doc,
        "benchmark solver에서 사용하는 single-relaxation-time collision은 distribution을 second-order equilibrium으로 relaxation시킨다.",
    )
    add_equation(doc, "fᵢ*(x) = fᵢ(x) - ω [ fᵢ(x) - fᵢᵉᵠ(ρ,u) ].")
    add_equation(
        doc,
        "fᵢᵉᵠ = wᵢρ [ 1 + (cᵢ·u)/cₛ² + (cᵢ·u)²/(2cₛ⁴) - (u·u)/(2cₛ²) ].",
        "D2Q9에서는 cₛ² = 1/3, ν = cₛ²(1/ω - 1/2)이다.",
    )
    add_para(
        doc,
        "collision 이후 streaming과 boundary treatment가 native LBM update operator에 의해 적용된다. 한 번의 collide-stream-boundary update를 G(f)라고 쓰면 정상상태 문제는 다음 fixed-point equation으로 표현된다.",
    )
    add_equation(doc, "R(f) = G(f) - f = 0.")
    add_para(
        doc,
        "여기서 R(f)를 native residual이라고 부르는 이유는 residual 평가가 baseline Picard LBM과 완전히 동일한 update, mask, wall rule, inlet/outlet treatment, obstacle geometry를 통과하기 때문이다. "
        "이 정의는 공정성 측면에서 중요하다. 모든 방법론은 동일한 discrete steady equation을 얼마나 빠르게 만족시키는지로 평가된다.",
    )

    add_heading(doc, "2.2 수렴 판정용 Macroscopic L2 Residual", 2)
    add_para(
        doc,
        "R(f)는 distribution space에서 정의되지만, 최종 목표는 안정적인 pressure-velocity steady field이므로 수렴 판정은 macroscopic variable을 기준으로 한다. "
        "연속된 accepted state fᵏ와 fᵏ⁺¹에 대해 pressure 및 velocity 변화량을 fluid cell에서만 계산한다.",
    )
    add_equation(
        doc,
        "rₚᵏ = ||pᵏ⁺¹ - pᵏ||₂ / √|Ω_f|,     rᵤᵏ = ||uₓᵏ⁺¹ - uₓᵏ||₂ / √|Ω_f|,     rᵥᵏ = ||uᵧᵏ⁺¹ - uᵧᵏ||₂ / √|Ω_f|.",
    )
    add_equation(
        doc,
        "r_macroᵏ = √[ (rₚᵏ)² + (rᵤᵏ)² + (rᵥᵏ)² + (r_wᵏ)² ].",
        "현재 2D benchmark에서는 z-velocity 항 r_w가 0이지만, 3D 확장을 위해 정의에는 포함한다.",
    )
    add_para(
        doc,
        "초기 상대 residual은 r_macroᵏ / r_macro⁰로 저장한다. 기존 f-RMS residual은 결과 파일에 보조 diagnostic으로 계속 저장하지만, 최종 protocol의 주 수렴 기준은 macroscopic L2 residual이다. "
        "이는 microscopic distribution의 변화가 작아 보여도 pressure 또는 velocity field가 아직 drift하는 경우를 피하기 위한 선택이다.",
    )

    add_heading(doc, "2.3 AP-Schur의 핵심 아이디어", 2)
    add_para(
        doc,
        "제안 solver는 native residual Jacobian의 Schur complement를 moment space에서 근사하는 방식에 기반한다. 정상상태 근방에서 Newton 방법은 다음 correction을 찾는다.",
    )
    add_equation(doc, "J_f(fᵏ)δf = -R(fᵏ),        J_f = ∂R/∂f.")
    add_para(
        doc,
        "하지만 복잡한 mask geometry와 LBM boundary operator를 포함한 전역 Jacobian J_f를 직접 구성하는 것은 비효율적이다. "
        "여기서 중요한 관찰은 정상상태 수렴 후반의 느린 error가 주로 hydrodynamic low-frequency moment mode에 존재한다는 점이다. 반면 kinetic mode는 collision과 native LBM smoothing에 의해 비교적 빠르게 감쇠한다. "
        "AP-Schur step은 residual equation을 compact moment space로 사영하고, hydrodynamic correction을 근사적으로 푼 뒤, 이를 distribution space로 lift한다.",
    )
    add_equation(
        doc,
        "m = P f,        Sₘ ≈ P J_f P†,        Sₘδm ≈ -P R(fᵏ),        δf_AP = P†δm.",
        "P는 distribution을 conserved/near-conserved moment로 보내는 projection이고, P†는 moment correction을 distribution으로 되돌리는 lifting operator이다.",
    )
    add_para(
        doc,
        "여기서 AP는 asymptotic-preserving의 의미로 사용된다. 즉, raw collide-stream relaxation만으로는 느리게 줄어드는 diffusive 또는 incompressible macroscopic mode에서도 preconditioner가 효과적으로 작동하도록 설계한다. "
        "Schur approximation은 standard Picard LBM이 느리게 제거하는 residual 성분을 직접 겨냥하며, detailed kinetic relaxation과 boundary consistency는 native LBM update가 유지한다.",
    )

    add_heading(doc, "2.4 Jacobian-Free Correction", 2)
    add_para(
        doc,
        "구현에서는 global Jacobian을 assemble하거나 저장하지 않기 위해 Jacobian-free product를 사용한다. trial direction v에 대해 native residual Jacobian action은 유한차분으로 근사된다.",
    )
    add_equation(doc, "J_f(f)v ≈ [ R(f + εv) - R(f) ] / ε,        ε = √ε_machine (1 + ||f||₂) / ||v||₂.")
    add_para(
        doc,
        "짧은 Krylov solve는 moment-preconditioned correction을 계산한다. AP-Schur operation은 느린 macroscopic residual 성분에 대한 left preconditioner 역할을 하며, residual evaluation 자체는 항상 원래의 LBM operator를 사용한다. "
        "따라서 memory footprint는 native solver와 크게 다르지 않고, 복잡한 mask와 boundary rule도 G(f)를 통해 일관되게 처리된다.",
    )

    add_heading(doc, "2.5 Candidate 생성과 Acceptance Rule", 2)
    add_para(
        doc,
        "각 outer cycle에서 solver는 보수적인 native LBM candidate와 AP-Schur candidate를 비교한다. AP-Schur correction은 무조건 적용되지 않으며, damping된 trial candidate를 순차적으로 평가한다.",
    )
    add_equation(doc, "f_trial(α) = fᵏ + αδf_AP,        α ∈ {1, 1/2, 1/4, ...}.")
    add_para(
        doc,
        "각 trial은 baseline solver와 동일한 LBM boundary 및 settling operation을 통과한 뒤 common macroscopic residual로 평가된다. "
        "trial이 residual을 감소시키고 물리적으로 admissible할 때에만 accept된다. 어떤 AP-Schur candidate도 acceptance gate를 통과하지 못하면 해당 cycle에서는 native LBM candidate를 사용한다. "
        "이 fallback 구조는 중요하다. AP-Schur는 안정적인 native operator를 대체하는 것이 아니라, 안전하게 residual을 낮출 수 있을 때만 개입하는 accelerator이다.",
    )
    add_simple_table(
        doc,
        ["Acceptance gate", "목적"],
        [
            ("Finite field check", "pressure, density, velocity, distribution에 NaN/Inf가 있으면 reject한다."),
            ("Positive density branch", "density가 비물리적 저밀도 branch로 이동하는 correction을 reject한다."),
            ("Residual monotonicity", "r_macro가 competing native candidate 또는 직전 accepted state보다 감소할 때만 accept한다."),
            ("Boundary consistency", "residual 평가 전에 native wall, inlet, outlet, mask operator를 반드시 다시 적용한다."),
            ("No reference-data injection", "analytic, Ghia, tight-reference 값은 solve 중 사용하지 않고 post-run error 계산에만 사용한다."),
        ],
        [2300, 7060],
    )

    add_heading(doc, "2.6 Geometry-Aware 처리", 2)
    add_para(
        doc,
        "검증 문제에는 rectangular channel, closed cavity, backward-facing step, cylinder mask, multi-cylinder array, T-junction open boundary가 포함된다. "
        "제안법은 이 모든 문제에 대해 하나의 geometry-aware AP-Schur-only algorithm을 사용한다. Geometry awareness는 세 가지 방식으로 들어간다. "
        "첫째, 모든 norm과 moment projection은 Ω_f에 대해서만 계산되어 solid node가 correction이나 residual에 섞이지 않는다. "
        "둘째, correction은 항상 native operator G(f)를 통해 평가되므로 bounce-back wall, inlet/outlet reconstruction, mask handling이 한 곳에서 일관되게 적용된다. "
        "셋째, open boundary 문제에서는 density와 velocity field가 물리적으로 의미 있는 branch에 있는지를 acceptance 단계에서 확인한다. "
        "이 safeguard는 Schur correction이 pressure gauge freedom이나 open-boundary 자유도를 이용해 residual만 낮추고 실제 물리장을 악화시키는 상황을 막는다.",
    )

    add_heading(doc, "2.7 수렴 조건과 Plateau Criterion", 2)
    add_para(
        doc,
        "수렴은 residual threshold와 residual-floor 조건을 모두 만족할 때만 인정한다. 먼저 absolute residual 기준은 benchmark tolerance τ에 대해 다음과 같이 둔다.",
    )
    add_equation(doc, "r_macroᵏ < 5τ.")
    add_para(
        doc,
        "추가로 AP-Schur에 의해 residual이 급격히 낮아진 직후 너무 빨리 종료되는 것을 막기 위해 plateau test를 사용한다. window size를 W라고 할 때 최근 window에서의 fractional improvement를 다음과 같이 정의한다.",
    )
    add_equation(doc, "I_Wᵏ = [ r_macroᵏ⁻ᵂ - r_macroᵏ ] / max(r_macroᵏ⁻ᵂ, ε_floor).")
    add_para(
        doc,
        "최근 window 동안 residual 감소율이 5% 이하이고 최소 LBM evaluation 수를 지난 경우 residual floor에 도달했다고 판단한다. "
        "즉 residual이 충분히 작아야 할 뿐 아니라, 현재 discrete problem에서 더 이상 의미 있게 줄지 않는 asymptotic floor에 가까워졌음을 확인한다. "
        "저장된 run이 residual threshold는 만족하지만 plateau를 만족하지 못하면 동일한 method와 동일한 convergence test로 저장 state에서 continuation을 수행한다. continuation은 iteration count만 늘리는 것이며 algorithm 변경이 아니다.",
    )

    add_heading(doc, "2.8 알고리즘", 2)
    add_simple_table(
        doc,
        ["Step", "Operation"],
        [
            ("1", "benchmark initial condition에서 f를 초기화하고 native boundary operator를 적용한다."),
            ("2", "p, u, v 및 초기 macroscopic L2 residual r_macro⁰를 fluid cell에서 계산한다."),
            ("3", "baseline Picard solver와 동일한 collide-stream-boundary operator G(f)로 conservative native LBM candidate를 생성한다."),
            ("4", "native residual을 moment space로 project하고, AP-Schur moment correction을 근사적으로 푼 뒤 distribution space로 lift한다."),
            ("5", "damped AP-Schur candidate를 시험한다. 각 trial마다 boundary rule을 재적용하고 common macroscopic residual을 계산한다."),
            ("6", "모든 admissibility gate를 통과하고 residual이 감소한 경우에만 AP-Schur candidate를 accept한다. 그렇지 않으면 native candidate를 accept한다."),
            ("7", "wall time, LBE calls, macro residual components, f-RMS diagnostic residual, AP-Schur trials/accepts를 기록한다."),
            ("8", "r_macro < 5τ 및 residual plateau/floor 조건을 모두 만족할 때만 종료한다. 아니면 latest accepted state에서 계속 진행한다."),
        ],
        [900, 8460],
    )

    add_heading(doc, "2.9 공정한 Benchmark Protocol", 2)
    add_para(
        doc,
        "제안 solver는 reference method와 동일한 residual definition, stopping rule, mesh scaling, post-processing metric으로 평가한다. "
        "최종 제안법은 AP-Schur-only이며 RRE는 비활성화한다. Ablation 결과 AP-Schur 단독 구성이 wall-clock 측면에서 가장 우수하면서 accuracy class를 유지했기 때문이다. "
        "특정 benchmark에만 들어가는 tuning coefficient는 사용하지 않는다. damping candidate, residual threshold, plateau tolerance, admissibility gate는 case-dependent knob이 아니라 global method setting이다.",
    )
    add_simple_table(
        doc,
        ["Benchmark family", "solve 이후 reference 비교"],
        [
            ("Channel Poiseuille", "analytic velocity profile 및 pressure-driven flow diagnostic."),
            ("Couette flow", "analytic linear velocity profile."),
            ("Lid-driven cavity Re100/Re400/Re1000", "Ghia et al. centerline velocity profile 및 field-level consistency metric."),
            ("Backward-facing step", "tight-reference 또는 충분히 수렴된 numerical field와 flow-feature diagnostic."),
            ("Cylinder wake Re40", "reference steady wake field 및 wake-related diagnostic."),
            ("Six-cylinder mask", "masked-domain velocity/pressure reference field."),
            ("T-junction", "open-boundary junction reference field 및 boundary-consistency diagnostic."),
        ],
        [2500, 6860],
    )

    add_heading(doc, "2.10 AP-Schur가 빠른 이유", 2)
    add_para(
        doc,
        "steady LBM Picard iteration은 각 step이 물리적 collide-stream update라서 robust하지만, 수렴 후반에 남는 long-wavelength hydrodynamic mode를 제거하는 데 느릴 수 있다. "
        "이 mode는 local relaxation만으로 약하게 감쇠되므로 global equilibration에 많은 lattice sweep이 필요하다. "
        "AP-Schur preconditioner는 이 병목을 직접 겨냥한다. 즉 macroscopic pressure-velocity Schur response를 근사하여 slow error가 존재하는 moment space에서 global correction을 제안한다. "
        "반면 kinetic 및 boundary-local error는 native update와 residual-monotone acceptance가 제어한다. 결과적으로 global hydrodynamic correction과 local LBM stability가 결합된다.",
    )
    add_equation(
        doc,
        "error = slow hydrodynamic moment component + fast kinetic component;     AP-Schur는 first component를 겨냥하고 native LBM은 second component를 감쇠시킨다.",
    )

    add_heading(doc, "2.11 재현성 및 논문 작성 시 강조점", 2)
    add_para(
        doc,
        "저장된 모든 history는 accepted macroscopic residual sequence, wall time, LBE-call count, final residual components, auxiliary f-RMS residual을 포함한다. "
        "현재 archive의 proposed row는 uniform_ap_schur_only로 기록되어 있으며, saved state에서 동일한 plateau criterion을 만족시키기 위해 continuation한 경우에만 continued label이 붙어 있다. "
        "이 label은 계산 경로를 기록하기 위한 것이며 수학적 방법 자체가 달라졌다는 뜻은 아니다.",
    )
    add_para(
        doc,
        "논문에서는 제안법을 native steady LBM residual에 대한 acceleration/preconditioning strategy로 제시하는 것이 가장 방어적이다. "
        "정확도 주장은 discretization, boundary treatment, Mach-number setting, mesh refinement와 함께 해석해야 한다. "
        "따라서 novelty claim은 AP-Schur가 물리적 steady equation을 바꾼다는 것이 아니라, geometry-aware, residual-safe, case-uniform Schur preconditioning을 통해 동일한 steady equation에 훨씬 빠르게 도달한다는 점에 두는 것이 적절하다.",
    )

    add_heading(doc, "2.12 Reviewer 대응 핵심 문장", 2)
    add_para(
        doc,
        "첫째, 제안법은 reference solution 또는 analytic/Ghia data를 solver 내부에 주입하지 않는다. 모든 reference 값은 계산이 끝난 뒤 error를 평가하기 위해서만 사용된다. "
        "둘째, AP-Schur-only solver는 case별 tuning coefficient를 사용하지 않는다. 동일한 correction 생성, damping candidate, acceptance gate, residual/plateau stopping rule이 모든 benchmark에 적용된다. "
        "셋째, AP-Schur correction은 residual-monotone admissibility test를 통과할 때에만 accepted되므로, 복잡한 geometry나 open boundary에서 비물리적 density/velocity branch로 이동하는 것을 방지한다. "
        "넷째, 최종 정확도는 동일한 LBM discretization과 boundary model에 의해 결정되며, AP-Schur는 그 discrete steady state로 가는 nonlinear iteration path를 가속하는 역할을 한다.",
    )

    add_heading(doc, "논문 인용 문헌 배치 제안", 2)
    add_para(
        doc,
        "최종 manuscript에는 standard LBM equilibrium/discretization 문헌, Ghia et al. lid-driven cavity benchmark, GMRES 및 Jacobian-free Newton-Krylov 문헌, Schur-complement/preconditioning 문헌을 함께 인용하는 것이 좋다. "
        "정확한 bibliography format은 최종 manuscript 조립 단계에서 journal style에 맞추어 정리하면 된다.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())

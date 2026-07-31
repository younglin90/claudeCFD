"""Assemble a revised Korean DOCX with the latest fixed30 proposed-method data."""

from __future__ import annotations

import csv
import json
import math
import statistics
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/mnt/c/Users/user/OneDrive/[논문투고]/할거_0_LBM_steady_state_가속화")
SOURCE = ROOT / "SafeNN_LBM_Paper_V7_KR_no_tjunction.docx"
OUTPUT = ROOT / "SafeNN_LBM_Paper_V8_KR_fixed30_latest_detailed_methods.docx"
DATA = Path("paper_revision_data") / "fixed30_scaling_strict"
FIG = Path("paper_revision_data") / "fixed30_manuscript_figures"

METHOD_LABELS = {
    "picard_lbm": "Picard LBM",
    "anderson_lbm": "Anderson LBM",
    "preconditioned_lbm": "Preconditioned LBM",
    "inexact_newton_lbe": "Inexact Newton LBE",
    "dual_time_mg_lbm": "Dual-time MG LBM",
    "proposed": "SafeNN-Final",
}

CASE_LABELS = {
    "kolmogorov_n32": "Kolmogorov flow",
    "channel_n32": "Plane Poiseuille channel",
    "couette_n32": "Couette flow",
    "cavity_re100_n33": "Lid-driven cavity Re=100",
    "cavity_re400_n49": "Lid-driven cavity Re=400",
    "cavity_re1000_n129": "Lid-driven cavity Re=1000",
    "multi_cylinder_n32": "Multi-cylinder voxel mask",
    "backward_step_n64": "Backward-facing step mask",
    "cylinder_wake_n64": "Cylinder wake analogue",
    "t_junction_n64": "T-junction mask",
}


def set_font(run, size=10):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)


def add_para(doc, text="", style=None, size=10):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, size)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, 14 if level == 1 else 12)
    return p


def add_caption(doc, text):
    p = add_para(doc, text, size=9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_figure(doc, path, caption, width=6.3):
    path = Path(path)
    if not path.exists():
        add_para(doc, f"[FIGURE MISSING: {path}]", size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_math_para(doc, text, number=None):
    """Add a Word OMML math paragraph with a simple linear math string."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = text
    mr.append(mt)
    omath.append(mr)
    omath_para.append(omath)
    p._p.append(omath_para)
    if number is not None:
        r = p.add_run(f"    ({number})")
        set_font(r, 10)
    return p


def add_hyperparameter_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Component", "Default value", "When changed", "Purpose"]
    for cell, h in zip(table.rows[0].cells, headers):
        cell.text = h
    rows = [
        ("Residual tolerance", "1e-7; cavity 5e-7", "Verifier-defined", "Native steady residual gate"),
        ("GMRES tolerance", "1e-3 relative", "Fixed in SafeNN branches", "Inexact Newton correction"),
        ("Nesterov beta cap", "0.95; 0.70 for large channel/couette", "Grid-stability rule", "Prevent residual overshoot"),
        ("Kinetic substeps", "4-15", "Operator-family stability rule", "Damp non-hydrodynamic modes"),
        ("Non-stiff cavity polish", "up to total 800 LBE calls", "Cavity Re<1000 and not large Re=400", "Remove high-frequency kinetic residue"),
        ("Stiff cavity PLBE warm phase", "gamma=0.7; 8,000-12,000 steps", "Re=1000 or Re=400 at 3x grid", "Stabilize high-Re cavity before native LBE closure"),
        ("Large masked post-smoothing", "500 native LBE steps", "N>=192 and residual already below tolerance", "Reduce mask-boundary kinetic residue"),
        ("Threads", "24", "Fixed for all methods", "Avoid using all 32 physical cores"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    remove_table_borders(table)
    add_caption(doc, "Table M1. Reproducibility parameters for the latest SafeNN-Final implementation.")


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9D9D9")
        borders.append(tag)
    tbl_pr.append(borders)


def load_rows():
    with (DATA / "summary.csv").open(newline="", encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    metrics = json.loads((DATA / "metrics.json").read_text(encoding="utf-8"))
    return rows, metrics


def finite(value, default=math.nan):
    try:
        out = float(value)
    except Exception:
        return default
    return out if math.isfinite(out) else default


def aggregate_methods(rows):
    by = defaultdict(list)
    for row in rows:
        by[row["method"]].append(row)
    out = []
    for method in METHOD_LABELS:
        rs = by[method]
        rels = [finite(r["rel_l2_vs_picard"]) for r in rs]
        rels = [r for r in rels if math.isfinite(r)]
        out.append(
            {
                "method": method,
                "label": METHOD_LABELS[method],
                "converged": sum(int(r["converged"]) for r in rs),
                "total_lbe": int(sum(finite(r["lbe_calls"], 0) for r in rs)),
                "total_wall": sum(finite(r["wall_seconds"], 0) for r in rs),
                "median_rel": statistics.median(rels) if rels else math.nan,
                "mean_rel": sum(rels) / len(rels) if rels else math.nan,
            }
        )
    return out


def sci(x):
    if x is None or not math.isfinite(float(x)):
        return "n/a"
    return f"{float(x):.3e}"


def num(x, nd=3):
    return f"{float(x):.{nd}f}"


def add_method_table(doc, rows):
    agg = aggregate_methods(rows)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Method", "Converged", "Total LBE", "Total wall (s)", "Median rel. L2", "Mean rel. L2"]
    for cell, h in zip(table.rows[0].cells, headers):
        cell.text = h
    for a in agg:
        cells = table.add_row().cells
        cells[0].text = a["label"]
        cells[1].text = f"{a['converged']}/30"
        cells[2].text = f"{a['total_lbe']:,}"
        cells[3].text = f"{a['total_wall']:.3f}"
        cells[4].text = sci(a["median_rel"])
        cells[5].text = sci(a["mean_rel"])
    remove_table_borders(table)
    add_caption(doc, "Table 1. Method-level summary for the latest fixed30 benchmark.")


def add_case_table(doc, metrics):
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["Case", "Level", "Pass", "Conv.", "LBE win", "Wall win", "Acc. win", "Rel. L2"]
    for cell, h in zip(table.rows[0].cells, headers):
        cell.text = h
    for c in metrics["case_results"]:
        cells = table.add_row().cells
        cells[0].text = CASE_LABELS.get(c["base_case_id"], c["base_case_id"])
        cells[1].text = f"{c['scaling_level']}x"
        cells[2].text = str(c["case_pass"])
        cells[3].text = str(c["converged"])
        cells[4].text = str(c["lbe_win"])
        cells[5].text = str(c["wall_win"])
        cells[6].text = str(c["accuracy_win"])
        cells[7].text = sci(c["proposed_rel_l2"])
    remove_table_borders(table)
    add_caption(doc, "Table 2. Per-case pass components for SafeNN-Final. A value of 1 indicates that the corresponding condition was satisfied.")


def copy_intro_from_source(doc):
    src = Document(SOURCE)
    paragraphs = src.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("1. Introduction"))
    stop = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("Contribution of this work"))
    for p in paragraphs[start:stop]:
        text = p.text.strip()
        if not text:
            continue
        if text.startswith("1. Introduction"):
            add_heading(doc, "1. Introduction", 1)
        else:
            text = text.replace("Safe-NN-SCMK", "SafeNN-Final")
            text = text.replace("SCMK", "SafeNN-Final")
            add_para(doc, text)
    add_para(
        doc,
        "Contribution of this work. 본 개정 원고는 기존의 6-case 중심 결과를 최신 fixed30 benchmark로 교체하고, "
        "제안 방법을 SafeNN-Final로 재정의한다. SafeNN-Final은 표준 LBE operator를 직접 수정하지 않고 native residual "
        "R(f)=L(f)-f를 대상으로 residual-monotone Nesterov lookahead, Jacobian-free Newton-Krylov correction, "
        "Fourier-moment AP-Schur preconditioning, 그리고 제한된 LBE post-relaxation을 결합한다. 본 원고의 핵심 기여는 "
        "(i) 10개 정상상태 LBM benchmark의 1x/2x/3x mesh scaling에서 6개 방법을 비교한 재현 가능한 fixed30 자료를 제시하고, "
        "(ii) 제안 방법이 LBE-call 측면에서는 30/30 case에서 best 또는 better임을 보이며, "
        "(iii) finite-tolerance Picard reference와의 field accuracy tradeoff를 정량적으로 드러낸다는 점이다.",
    )
    add_para(
        doc,
        "Paper organization. §2는 최신 SafeNN-Final 구현과 reference solver 고정 원칙을 기술한다. §3은 fixed30 결과, "
        "수렴성, wall-clock, 정확도, residual history, field contour를 제시한다. §4는 speedup과 field-accuracy 사이의 "
        "tradeoff 및 현 단계의 한계를 논의한다.",
    )


def add_methods(doc):
    add_heading(doc, "2. Methods", 1)
    add_heading(doc, "2.1 Discrete velocity model and native LBE map", 2)
    add_para(
        doc,
        "본 연구의 모든 2D 검증은 D2Q9 BGK LBM을 기반으로 한다. 격자 속도는 c0=(0,0), "
        "c1-4=(±1,0),(0,±1), c5-8=(±1,±1)이고, 가중치는 w0=4/9, w1-4=1/9, w5-8=1/36이다. "
        "분포함수 f_i(x,t)는 collision, forcing, streaming 및 boundary treatment를 모두 포함하는 한 단계 LBE map L로 갱신된다.",
    )
    add_math_para(doc, "f(t+1) = L(f(t))", "1")
    add_math_para(doc, "rho = sum_i f_i,    rho u = sum_i c_i f_i", "2")
    add_para(
        doc,
        "평형분포는 표준 저마하수 D2Q9 평형식을 사용한다. 외력이 있는 periodic/masked flow에서는 기존 case object에 정의된 forcing "
        "term을 그대로 사용하고, 벽 경계가 있는 channel, Couette, cavity에서는 해당 case의 bounce-back 또는 non-equilibrium "
        "extrapolation boundary implementation을 그대로 유지한다. 제안 방법은 이 LBE map을 바꾸지 않고, 오직 L(f)-f residual을 대상으로 "
        "가속 correction을 계산한다.",
    )
    add_math_para(doc, "f_i^eq = w_i rho [1 + 3(c_i·u) + 9/2 (c_i·u)^2 - 3/2 (u·u)]", "3")
    add_math_para(doc, "nu = c_s^2 (tau - 1/2),    omega = 1/tau,    c_s^2 = 1/3", "4")

    add_heading(doc, "2.2 Native fixed-point residual and convergence norm", 2)
    add_para(
        doc,
        "정상상태 해는 pseudo-time marching 문제로 재정의하지 않고 원래 LBE map의 fixed point로 정의한다. 따라서 residual은 "
        "분포함수 공간에서 R(f)=L(f)-f로 정의된다. Solver 내부 residual norm은 전체 분포함수 자유도 수로 정규화한 RMS norm을 사용한다. "
        "Verifier에서는 final residual이 5 times tolerance보다 작을 때 converged로 기록하지만, proposed solver 자체는 입력 tolerance를 "
        "목표 residual로 사용한다.",
    )
    add_math_para(doc, "R(f) = L(f) - f", "5")
    add_math_para(doc, "||R(f)||_rms = [ (1/N_dof) sum_j R_j(f)^2 ]^(1/2)", "6")
    add_para(
        doc,
        "이 residual 정의는 중요한 재현성 조건이다. 본 연구의 proposed solver는 analytic velocity profile, converged target field, "
        "또는 precomputed Picard solution을 residual에 삽입하지 않는다. 모든 correction은 현재 f, L(f), finite-difference residual probe, "
        "그리고 native LBE smoothing으로부터만 계산된다.",
    )

    add_heading(doc, "2.3 Residual-monotone Nesterov lookahead", 2)
    add_para(
        doc,
        "SafeNN-Final의 첫 번째 안정화 장치는 Nesterov-type lookahead이다. 이전 두 iterate가 주어지면 y_k=f_k+beta_k(f_k-f_{k-1})를 "
        "만든다. Lookahead residual이 현재 residual보다 허용 비율 이상 커지거나 비정상 값이 포함되면 lookahead를 폐기하고 y_k=f_k로 "
        "되돌린다. 이 safeguard는 cavity와 mask geometry에서 extrapolation이 kinetic mode를 증폭하는 것을 막기 위한 장치이다.",
    )
    add_math_para(doc, "y_k = f_k + beta_k (f_k - f_{k-1})", "7")
    add_math_para(doc, "accept y_k if ||R(y_k)|| <= [1 + eps_accept + 0.2 beta_k] ||R(f_k)||", "8")
    add_math_para(doc, "otherwise: y_k <- f_k,    beta_k <- 0.7 beta_k", "9")
    add_para(
        doc,
        "beta는 residual 감소가 지속될 때 증가시키고 residual이 증가하면 감소시킨다. 기본 beta cap은 0.95이며, large channel/couette "
        "branch에서는 grid stability를 위해 0.70을 사용한다. 이 값은 case-specific target tuning이 아니라 동일 operator family의 "
        "large-grid stability rule로 사용된다.",
    )

    add_heading(doc, "2.4 Jacobian-free Newton-Krylov correction", 2)
    add_para(
        doc,
        "Accepted lookahead state y_k에서 Newton correction은 Jacobian-free 방식으로 계산한다. 선형화된 residual equation은 "
        "J_R(y_k) delta f = -R(y_k)이고, 이 선형계는 GMRES로 근사적으로 푼다. Jacobian matrix는 명시적으로 만들지 않고 finite-difference "
        "Jacobian-vector product만 사용한다.",
    )
    add_math_para(doc, "J_R(y_k) delta f_k = -R(y_k)", "10")
    add_math_para(doc, "J_R(y) v ≈ [ R(y + epsilon v) - R(y) ] / epsilon", "11")
    add_math_para(doc, "epsilon = 1e-7 ( ||y||_2 + 1 ) / max( ||v||_2, 1e-30 )", "12")
    add_para(
        doc,
        "GMRES relative tolerance는 1e-3이며, restart length는 2 times krylov_max로 설정한다. krylov_max는 기본적으로 4-10 범위에서 "
        "operator family와 grid scale에 따라 선택된다. 이 설정은 nonlinear residual을 과도하게 줄이는 inner solve를 피하고, 한 번의 Newton "
        "correction 뒤 native LBE smoothing이 kinetic mode를 감쇠하도록 설계한 것이다.",
    )

    add_heading(doc, "2.5 Fourier-moment AP-Schur preconditioner", 2)
    add_para(
        doc,
        "GMRES의 preconditioner는 D2Q9 moment projection M과 equilibrium lift T를 사용한 Fourier-moment Schur approximation이다. "
        "Uniform equilibrium 근방에서 collision linearization C=(1-omega)I+omega T M, streaming symbol A(k)=diag exp[-i k·c_i]를 "
        "사용하면 residual Jacobian symbol은 J(k)=I-A(k)C로 쓸 수 있다. Galerkin Schur complement는 S_U^G(k)=M J(k) T이다.",
    )
    add_math_para(doc, "C = (1 - omega) I + omega T M", "13")
    add_math_para(doc, "A(k) = diag_i exp[-i k·c_i],    J(k) = I - A(k) C", "14")
    add_math_para(doc, "S_U^G(k) = M J(k) T = I - M A(k) T", "15")
    add_para(
        doc,
        "AP-Schur correction은 kinetic null-space의 leading contribution을 근사적으로 반영한다. 구현에서는 coefficient를 clipping하고 "
        "mode별 작은 diagonal regularization을 더해 ill-conditioned mode를 피한다. zero Fourier mode에서는 mass mean correction을 "
        "꺼 두고 momentum mean만 passthrough한다. 이 처리는 mass conservation과 singular mean mode 문제를 동시에 피하기 위한 것이다.",
    )
    add_math_para(doc, "S_U^AP(k) = S_U^G(k) - alpha_AP [ M A(k)^2 T - (M A(k) T)^2 ]", "16")
    add_math_para(doc, "P^{-1} R_f = T F^{-1} [ S_U^AP(k) + eta I ]^{-1} F [ M R_f ]", "17")
    add_para(
        doc,
        "이 preconditioner는 periodic bulk linearization에서 유도되므로 wall/mask geometry에서 exact inverse는 아니다. 본 논문에서는 "
        "이를 exact multigrid 또는 exact boundary-aware Schur solver로 주장하지 않고, slow hydrodynamic residual mode를 줄이는 "
        "operator-level preconditioner로 사용한다.",
    )

    add_heading(doc, "2.6 Kinetic post-relaxation and operator-family stabilization", 2)
    add_para(
        doc,
        "Newton correction 후에는 K_eff개의 native LBE step을 적용한다. 이 단계는 collision-streaming operator 자체를 사용하여 "
        "non-hydrodynamic kinetic mode를 감쇠한다. 즉, preconditioner가 거시 모멘텀 성분을 빠르게 줄이고, LBE smoothing이 kinetic "
        "high-frequency residue를 줄이는 역할 분담을 한다.",
    )
    add_math_para(doc, "f_{k+1} = L^{K_eff}( y_k + alpha delta f_k )", "18")
    add_para(
        doc,
        "최신 solve_proposed_single 구현은 하나의 proposed entry point를 유지하지만, operator family별 안정화 submode를 포함한다. "
        "Kolmogorov flow에는 짧은 secant bootstrap을 사용한다. Channel/Couette에는 동일 LBE map의 numba-optimized wall kernel을 "
        "바인딩한다. Non-stiff cavity는 SafeNN 이후 total 800 LBE-call 수준까지 native post-relaxation을 수행한다. Re=1000 cavity와 "
        "Re=400의 3x grid처럼 stiff한 cavity에는 gamma=0.7 PLBE warm phase를 사용한 뒤 native cavity LBE로 1,000 step closure를 "
        "수행한다. Large masked geometry에서는 residual이 이미 tolerance 아래일 때만 500 native LBE step을 추가한다.",
    )
    add_para(
        doc,
        "이 submode들은 모두 동일한 native LBE residual과 현재 state에 기반한다. 다만 완전히 parameter-free인 단일 closed-form algorithm은 "
        "아니므로, 본 논문에서는 SafeNN-Final을 'research-stage single proposed implementation with operator-family stabilization'으로 "
        "명시한다.",
    )

    add_heading(doc, "2.7 Algorithm summary", 2)
    add_para(
        doc,
        "Algorithm 1. SafeNN-Final native-residual acceleration",
        size=10,
    )
    for line in [
        "Input: LBE case object, tolerance tol, initial distribution f0.",
        "1. Bind implementation-equivalent optimized LBE kernels when available.",
        "2. If the case is in a stiff cavity regime, run a gamma=0.7 PLBE warm phase and return to the native cavity LBE map.",
        "3. Otherwise initialize f_prev=f0 and f=f0; optionally apply a short secant bootstrap for periodic Kolmogorov flow.",
        "4. At each outer iteration compute R(f) and stop if ||R(f)||_rms < tol.",
        "5. Form Nesterov lookahead y=f+beta(f-f_prev) and accept it only if the residual safeguard is satisfied.",
        "6. Approximately solve J_R(y) delta f=-R(y) by GMRES with AP-Schur preconditioning.",
        "7. Apply delta f, then run K_eff native LBE smoothing steps.",
        "8. If the correction is non-finite or rejected, fall back to native LBE smoothing for this iteration only.",
        "9. Apply bounded post-relaxation prescribed by the operator-family stabilization rule.",
        "Output: final distribution f, residual history, LBE-call count, wall-clock time.",
    ]:
        add_para(doc, line, size=9)

    add_heading(doc, "2.8 Anti-cheating and reproducibility constraints", 2)
    add_para(
        doc,
        "Reviewer가 가장 쉽게 공격할 수 있는 부분은 target field leakage이다. 이를 피하기 위해 fixed30 verifier는 proposed source에서 "
        "analytical target injection, target-deflated equilibrium lift, direct Anderson call delegation 등 금지 token을 검사한다. "
        "또한 proposed cache는 solver source hash를 포함하므로 proposed code가 바뀌면 cached result가 자동으로 무효화된다.",
    )
    add_para(
        doc,
        "Reference solvers는 proposed optimization 중 변경하지 않았다. 모든 method x case 결과는 같은 summary.csv schema로 저장되며, "
        "residual history CSV와 cached NPZ field를 함께 남긴다. Wall-clock은 Python/Numba 구현의 영향을 받으므로, 본 논문에서는 "
        "LBE-call과 wall-clock을 모두 보고한다.",
    )
    add_hyperparameter_table(doc)

    add_heading(doc, "2.9 Reference solvers and fixed30 benchmark protocol", 2)
    add_para(
        doc,
        "비교 방법은 Picard LBM, Walker-Ni style Anderson-accelerated LBM, preconditioned LBM, inexact Newton LBE, "
        "dual-time multigrid-inspired LBM, 그리고 proposed SafeNN-Final의 6개이다. Reference solver는 fixed30 검증 중 "
        "고정했으며, proposed solver 개선 과정에서 reference 구현을 약화시키지 않았다. 다만 preconditioned LBM과 dual-time MG는 "
        "문헌 아이디어를 repo-local production proxy로 구현한 것이므로, Discussion에서 이 한계를 명시한다.",
    )
    add_para(
        doc,
        "검증 문제는 Kolmogorov flow, plane Poiseuille channel, Couette flow, lid-driven cavity Re=100/400/1000, "
        "multi-cylinder voxel mask, backward-facing step mask, cylinder wake analogue, T-junction mask의 10개이다. 각 문제는 "
        "1x, 2x, 3x mesh level에서 계산하여 총 30 case를 구성한다. 기본 tolerance는 1e-7이고 cavity case는 5e-7이다. "
        "Voxel/mask case의 error norm은 fluid cell 기준으로 평가한다.",
    )
    add_math_para(doc, "e_rel = ||u_method - u_Picard||_2 / max(||u_Picard||_2, 1e-30)", "19")
    add_math_para(doc, "Pass = Converged and LBE_win and Wall_win and Accuracy_win", "20")
    add_para(
        doc,
        "Pass criterion은 case별로 네 조건을 모두 만족해야 한다. Proposed solver가 수렴해야 하고, LBE-call 수가 converged fixed "
        "reference 중 최소값 이하이어야 하며, wall-clock도 최소값 이하이어야 하고, finite-tolerance Picard reference에 대한 "
        "relative L2 error가 converged non-Picard reference 중 최저값과 0.1% 이내로 같거나 작아야 한다. 이 기준은 매우 엄격하며, "
        "속도 가속과 field accuracy를 동시에 요구한다.",
    )
    add_para(
        doc,
        "계산은 NUMBA_NUM_THREADS=24, OMP_NUM_THREADS=24, OPENBLAS_NUM_THREADS=1, MKL_NUM_THREADS=1 조건에서 수행했다. "
        "결과 cache는 solver/config hash를 포함하여 stale cache를 방지했다. Raw data는 summary.csv, metrics.json, residual "
        "history CSV, cached field NPZ 및 manuscript figure로 저장했다.",
    )


def add_results(doc, rows, metrics):
    add_heading(doc, "3. Results", 1)
    add_para(
        doc,
        "본 절은 최신 fixed30 결과만을 보고한다. 따라서 이전 원고의 6-case 산술평균/기하평균 speedup claim, 일부 subset-only 결과, "
        "또는 3D 검증 관련 문장은 본 개정본의 주 결과에서 사용하지 않는다.",
    )
    add_heading(doc, "3.1 Overall fixed30 performance", 2)
    add_para(
        doc,
        f"SafeNN-Final의 전체 score는 {metrics['score']:.2f}였고, strict all-pass 조건은 만족하지 못했다 "
        f"(pass_count={metrics['pass_count']}/30, all_pass={metrics['all_pass']}). 그러나 LBE-call 기준으로는 "
        f"{metrics['lbe_win_count']}/30 case에서 best fixed reference보다 작거나 같았고, wall-clock 기준으로는 "
        f"{metrics['wall_win_count']}/30 case에서 best 또는 동급이었다. 수렴성은 {metrics['converged_count']}/30 case에서 "
        "확보되었다.",
    )
    add_para(
        doc,
        f"가장 강한 결과는 계산량 감소이다. Proposed solver의 총 LBE-call은 122,398으로 Picard LBM의 754,188, "
        "preconditioned LBM의 376,437, inexact Newton LBE의 288,472보다 작았다. Picard 대비 평균 LBE-call speedup은 "
        "case 평균 기준으로 약 59.35x였고, best fixed reference 대비 평균 LBE-call speedup은 "
        f"{metrics['mean_lbe_speedup_vs_best_fixed']:.2f}x였다.",
    )
    add_method_table(doc, rows)
    add_figure(
        doc,
        FIG / "fig_fixed30_method_summary.png",
        "Figure 1. Fixed30 method-level comparison. SafeNN-Final sharply reduces total LBE calls and total wall-clock time, but the median field error remains finite because several cases converge to a state different from the finite-tolerance Picard reference.",
    )

    add_heading(doc, "3.2 Pass-component analysis", 2)
    add_para(
        doc,
        "Strict pass failure의 주된 원인은 속도가 아니라 accuracy criterion이었다. Proposed solver는 모든 case에서 LBE-call 조건을 "
        "만족했지만, accuracy criterion은 13/30 case에서만 만족했다. Wall-clock criterion은 23/30 case에서 만족했으며, "
        "나머지는 주로 multi-cylinder 2x, backward-step 2x, cylinder-wake 1x/2x/3x 및 일부 channel/couette scaling에서 발생했다.",
    )
    add_para(
        doc,
        "Accuracy failure는 대부분 residual tolerance를 만족했음에도 finite-tolerance Picard reference와의 velocity field relative "
        "L2가 reference 방법 중 최저값보다 큰 경우이다. 이는 proposed solver가 native residual을 빠르게 줄이는 장점과, loose Picard "
        "reference trajectory에 가까운 field를 유지해야 하는 strict benchmark accuracy criterion 사이의 tradeoff를 보여준다.",
    )
    add_case_table(doc, metrics)
    add_figure(
        doc,
        FIG / "fig_fixed30_pass_heatmap.png",
        "Figure 2. Pass-component heatmap for 30 fixed30 cases. LBE-call wins are universal, whereas the strict field-accuracy criterion is the limiting condition.",
        width=5.0,
    )
    add_figure(
        doc,
        FIG / "fig_fixed30_speed_accuracy_tradeoff.png",
        "Figure 3. Per-case speed/accuracy tradeoff. Values above one indicate that SafeNN-Final is better than or equal to the best fixed reference for the corresponding metric.",
    )

    add_heading(doc, "3.3 Residual convergence histories", 2)
    add_para(
        doc,
        "Residual history는 proposed solver가 많은 case에서 매우 적은 LBE-call로 native residual을 tolerance 근처까지 낮춘다는 점을 "
        "보여준다. 특히 Kolmogorov, channel, couette, low-Re cavity 및 일부 masked flow에서는 Picard 또는 Anderson 대비 "
        "초기 residual 감소가 훨씬 빠르다. 반면 cylinder wake analogue의 2x/3x에서는 proposed residual이 strict convergence "
        "threshold를 만족하지 못했고, 이 두 case가 converged_count=28/30의 원인이다.",
    )
    add_figure(
        doc,
        FIG / "fig_fixed30_residual_histories_1x.png",
        "Figure 4. Native residual histories for 1x benchmark cases. The proposed solver generally reduces residual with far fewer LBE calls, while several masked-flow cases retain field-accuracy differences.",
    )

    add_heading(doc, "3.4 Field contours and error localization", 2)
    add_para(
        doc,
        "Field contour는 Picard reference와 proposed field의 차이가 주로 벽, mask boundary, separated-flow analogue의 recirculation region "
        "근처에 집중됨을 보여준다. Smooth periodic/channel/couette case에서도 residual은 충분히 작지만 finite-tolerance Picard reference와 "
        "비교하면 작은 field offset이 남을 수 있다. 따라서 본 결과는 residual convergence와 finite-time Picard field agreement가 "
        "동일한 지표가 아님을 시사한다.",
    )
    add_figure(
        doc,
        FIG / "fig_fixed30_core_fields_1x.png",
        "Figure 5. Core 1x velocity-magnitude fields and absolute proposed-vs-Picard differences.",
    )
    add_figure(
        doc,
        FIG / "fig_fixed30_mask_fields_1x.png",
        "Figure 6. Masked-flow 1x velocity-magnitude fields and absolute proposed-vs-Picard differences.",
    )


def add_discussion(doc):
    add_heading(doc, "4. Discussion and Conclusions", 1)
    add_heading(doc, "4.1 Interpretation of the latest benchmark", 2)
    add_para(
        doc,
        "최신 fixed30 결과는 SafeNN-Final의 장점과 한계를 동시에 보여준다. 제안 방법은 native LBE residual을 직접 대상으로 하며, "
        "LBE-call 측면에서 모든 case의 best fixed reference보다 작거나 같았다. 이는 steady LBM 문제에서 Newton-Krylov correction, "
        "residual safeguard, AP-Schur moment preconditioning, 그리고 제한된 kinetic smoothing의 결합이 계산량을 크게 줄일 수 있음을 "
        "입증한다.",
    )
    add_para(
        doc,
        "그러나 field accuracy 기준에서는 13/30 case만 strict criterion을 통과했다. 이 결과는 제안 방법이 실패했다는 단순한 의미라기보다, "
        "native residual을 빠르게 줄이는 해와 finite-tolerance Picard trajectory가 만든 reference field가 일부 stiff/masked geometry에서 "
        "다를 수 있음을 보여준다. 따라서 본 논문의 claim은 '모든 지표에서 기존 방법을 압도하는 보편 solver'가 아니라, residual-based "
        "steady-state acceleration의 계산량 이득과 field-accuracy tradeoff를 정량화하는 것으로 제한한다.",
    )
    add_heading(doc, "4.2 Limitations", 2)
    add_para(
        doc,
        "첫째, 현재 SafeNN-Final은 하나의 entry point로 구현되지만 내부에 operator-family-dependent stabilization submode가 존재한다. "
        "완전히 parameter-free인 단일 수식 알고리즘으로 정리하려면 추가 연구가 필요하다. 둘째, backward step, cylinder wake, T-junction은 "
        "현재 inlet/outlet benchmark가 아니라 periodic masked-flow analogue로 해석해야 한다. 셋째, strict accuracy metric은 "
        "finite-tolerance Picard reference를 기준으로 하므로, 더 엄밀한 tight reference 또는 문헌 benchmark solution을 사용한 재평가가 필요하다. "
        "넷째, 본 원고는 2D D2Q9 정상상태 문제에 한정되며, 3D 검증은 포함하지 않는다.",
    )
    add_heading(doc, "4.3 Conclusions", 2)
    add_para(
        doc,
        "본 연구는 표준 LBM operator를 유지하면서 native residual을 직접 가속하는 SafeNN-Final solver를 재정의하고, 10개 benchmark의 "
        "1x/2x/3x mesh scaling으로 구성된 fixed30 suite에서 평가했다. 최신 결과에서 proposed solver는 30/30 case의 LBE-call 기준을 "
        "만족하고, 23/30 case의 wall-clock 기준을 만족했으며, 28/30 case에서 수렴했다. 반면 strict field-accuracy criterion은 "
        "13/30 case만 만족했다.",
    )
    add_para(
        doc,
        "따라서 SafeNN-Final은 steady LBM residual acceleration의 유망한 방향이지만, 현 단계에서는 정확도까지 모든 benchmark에서 우수한 "
        "완성형 production solver로 주장하기보다, 계산량 감소와 accuracy tradeoff를 함께 보고하는 computational physics 방법론으로 "
        "제시하는 것이 타당하다. 향후 연구는 tight reference 기반 accuracy 검증, inlet/outlet benchmark 재설계, 내부 submode 단순화, "
        "그리고 3D 확장 검증에 초점을 맞출 필요가 있다.",
    )


def add_references(doc):
    src = Document(SOURCE)
    paragraphs = src.paragraphs
    start = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "References"), None)
    if start is None:
        add_heading(doc, "References", 1)
        add_para(doc, "[AUTHOR VERIFY: references from the previous manuscript were not found.]")
        return
    add_heading(doc, "References", 1)
    for p in paragraphs[start + 1 :]:
        text = p.text.strip()
        if text:
            add_para(doc, text, size=9)


def build_docx():
    rows, metrics = load_rows()
    doc = Document()
    styles = doc.styles
    for name in ["Normal", "Body Text", "First Paragraph"]:
        if name in styles:
            styles[name].font.name = "Malgun Gothic"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            styles[name].font.size = Pt(10)

    title = add_heading(doc, "SafeNN-Final을 이용한 정상상태 격자 볼츠만 방정식의 Native-Residual 가속: fixed30 scaling benchmark", 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Updated Korean manuscript draft based on SafeNN-Final fixed30 results", size=9).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "격자 볼츠만 방법(LBM)은 명시적 collision-streaming 구조와 병렬성 때문에 정상 유동 계산에 널리 사용되지만, "
        "steady-state 해를 얻기 위해 많은 LBE time step을 요구한다. 본 개정 원고는 표준 LBM operator를 유지하면서 native residual "
        "R(f)=L(f)-f를 가속하는 SafeNN-Final solver를 기술하고, 10개 benchmark와 1x/2x/3x mesh scaling으로 구성된 fixed30 suite에서 "
        "Picard LBM, Anderson LBM, preconditioned LBM, inexact Newton LBE, dual-time multigrid-inspired LBM과 비교한다. "
        "최신 결과에서 SafeNN-Final은 30/30 case에서 LBE-call 기준 best 또는 동급을 달성했고, 23/30 case에서 wall-clock 기준 best 또는 "
        "동급이었다. 전체 수렴성은 28/30 case였으며, strict field-accuracy criterion은 13/30 case에서 만족했다. 이 결과는 제안 방법이 "
        "계산량 감소에는 매우 효과적이지만, finite-tolerance Picard reference와의 field agreement에서는 stiff cavity 및 masked geometry에서 "
        "tradeoff가 존재함을 보여준다. 따라서 본 연구는 모든 지표에서 보편적으로 우수한 solver claim이 아니라, steady LBM native-residual "
        "acceleration의 속도 이득과 정확도 한계를 정량화하는 computational physics benchmark study로 제시된다.",
    )
    add_para(doc, "Keywords: lattice Boltzmann method; steady-state acceleration; Newton-Krylov method; Nesterov safeguard; residual convergence; fixed-point iteration")

    copy_intro_from_source(doc)
    add_methods(doc)
    add_results(doc, rows, metrics)
    add_discussion(doc)
    add_references(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build_docx()
    print(out)

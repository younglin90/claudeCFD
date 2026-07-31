#!/usr/bin/env python3
"""Build the JCP-style draft .docx: Times New Roman 10 pt reference styles + pandoc (native OMML equations)."""
import subprocess, sys
from pathlib import Path
import pypandoc
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
MD = HERE / "paper_draft.md"
REF = HERE / "_reference.docx"
OUT = HERE / "electrospray_solver_JCP_draft.docx"

# 1) default pandoc reference doc
data = subprocess.run([pypandoc.get_pandoc_path(), "--print-default-data-file", "reference.docx"],
                      capture_output=True, check=True).stdout
REF.write_bytes(data)

# 2) restyle: Times New Roman, 10 pt body, black headings
doc = Document(str(REF))
BLACK = RGBColor(0, 0, 0)

def set_font(style, size, bold=None, italic=None, color=BLACK, name="Times New Roman"):
    f = style.font
    f.name = name
    f.size = Pt(size)
    if bold is not None: f.bold = bold
    if italic is not None: f.italic = italic
    if color is not None: f.color.rgb = color
    # east-asian font binding so TNR applies everywhere
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        from docx.oxml.ns import qn
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    from docx.oxml.ns import qn
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), name)
    # theme font bindings override explicit names in Word/LO — strip them
    for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme'):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]

styles = doc.styles
plan = {
    "Normal": dict(size=10),
    "Body Text": dict(size=10),
    "First Paragraph": dict(size=10),
    "Compact": dict(size=9),           # table cells slightly smaller, JCP-like
    "Title": dict(size=14, bold=True),
    "Subtitle": dict(size=11, italic=True),
    "Author": dict(size=11),
    "Date": dict(size=10),
    "Abstract": dict(size=9.5),
    "Heading 1": dict(size=12, bold=True),
    "Heading 2": dict(size=11, bold=True),
    "Heading 3": dict(size=10, bold=True, italic=True),
    "Heading 4": dict(size=10, bold=False, italic=True),
    "Block Text": dict(size=10),
    "Caption": dict(size=9),
    "Table Caption": dict(size=9),
    "Image Caption": dict(size=9),
}
for name, kw in plan.items():
    try:
        set_font(styles[name], **kw)
    except KeyError:
        pass

# body justified, modest spacing
for name in ("Normal", "Body Text", "First Paragraph"):
    try:
        pf = styles[name].paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(4)
        pf.line_spacing = 1.08
    except KeyError:
        pass
# headings: no page-break-before surprises, black text guaranteed above
for name in ("Heading 1", "Heading 2", "Heading 3"):
    try:
        pf = styles[name].paragraph_format
        pf.space_before = Pt(10); pf.space_after = Pt(4)
    except KeyError:
        pass
# Title centered
try:
    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Author"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
except KeyError:
    pass
doc.save(str(REF))

# 3) convert (math -> native Word OMML equations)
pypandoc.convert_file(str(MD), "docx", outputfile=str(OUT),
                      extra_args=["--reference-doc=" + str(REF), "--standalone"])

# 4) sanity: reopen, count paragraphs/equations/images
out = Document(str(OUT))
import re
xml = out.element.xml
n_eq = xml.count("oMath")
n_img = xml.count("graphicData")
print(f"OK: {OUT}")
print(f"paragraphs={len(out.paragraphs)} tables={len(out.tables)} omml_math_nodes={n_eq} images~={n_img}")
